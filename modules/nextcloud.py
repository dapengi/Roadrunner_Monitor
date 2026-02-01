"""
Nextcloud integration functions
"""

import requests
import os
import datetime
import re
import logging
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from config import (
    NEXTCLOUD_URL, NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN,
    NEXTCLOUD_BASE_FOLDER, DOWNLOAD_DIR, N8N_WEBHOOK_URL
)
from data.committee_mapping import parse_meeting_title

logger = logging.getLogger(__name__)


def send_n8n_webhook(filename, folder_path):
    """Send webhook notification to n8n after successful file upload."""
    try:
        if not N8N_WEBHOOK_URL:
            logger.warning("N8N webhook URL not configured, skipping webhook notification")
            return False

        payload = {
            "filename": filename,
            "folder_path": folder_path,
        }

        response = requests.post(
            N8N_WEBHOOK_URL,
            json=payload,
            timeout=10
        )

        if response.status_code == 200:
            logger.info(f"Successfully sent webhook notification to n8n for {filename}")
            return True
        else:
            logger.warning(f"Webhook notification failed: {response.status_code} - {response.text}")
            return False

    except Exception as e:
        logger.error(f"Error sending webhook notification: {e}")
        return False


def create_nextcloud_folder(folder_path):
    """Create a folder in Nextcloud using WebDAV."""
    try:
        webdav_url = f"{NEXTCLOUD_URL}/remote.php/dav/files/{NEXTCLOUD_USERNAME}/{folder_path}"
        
        response = requests.request(
            'MKCOL',
            webdav_url,
            auth=(NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN),
            timeout=30
        )
        
        # 201 = created, 405 = already exists
        if response.status_code in [201, 405]:
            logger.info(f"Nextcloud folder ready: {folder_path}")
            return True
        else:
            logger.error(f"Failed to create Nextcloud folder: {response.status_code} - {response.text}")
            return False
            
    except Exception as e:
        logger.error(f"Error creating Nextcloud folder: {e}")
        return False


def upload_to_nextcloud(file_path, nextcloud_path):
    """Upload a file to Nextcloud using WebDAV."""
    try:
        webdav_url = f"{NEXTCLOUD_URL}/remote.php/dav/files/{NEXTCLOUD_USERNAME}/{nextcloud_path}"
        
        with open(file_path, 'rb') as f:
            response = requests.put(
                webdav_url,
                data=f,
                auth=(NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN),
                timeout=300  # 5 minutes for file upload
            )
        
        if response.status_code in [201, 204]:
            logger.info(f"Successfully uploaded to Nextcloud: {nextcloud_path}")
            return True
        else:
            logger.error(f"Failed to upload to Nextcloud: {response.status_code} - {response.text}")
            return False
            
    except Exception as e:
        logger.error(f"Error uploading to Nextcloud: {e}")
        return False


def get_nextcloud_share_link(nextcloud_path):
    """Create a public share link for a Nextcloud file with edit permissions."""
    try:
        # Create share via OCS API
        share_url = f"{NEXTCLOUD_URL}/ocs/v2.php/apps/files_sharing/api/v1/shares"
        
        data = {
            'path': f"/{nextcloud_path}",
            'shareType': 3,  # Public link
            'permissions': 15  # Full permissions (read + write + create + delete + share)
        }
        
        headers = {
            'OCS-APIRequest': 'true',
            'Content-Type': 'application/x-www-form-urlencoded'
        }
        
        response = requests.post(
            share_url,
            data=data,
            headers=headers,
            auth=(NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN),
            timeout=30
        )
        
        if response.status_code == 200:
            # Parse XML response to get share URL
            from xml.etree import ElementTree as ET
            root = ET.fromstring(response.text)
            
            url_element = root.find('.//url')
            if url_element is not None:
                share_link = url_element.text
                logger.info(f"Created Nextcloud editable share link: {share_link}")
                return share_link
            else:
                logger.error("Could not find URL in share response")
                return None
        else:
            logger.error(f"Failed to create share: {response.status_code} - {response.text}")
            return None
            
    except Exception as e:
        logger.error(f"Error creating Nextcloud share link: {e}")
        return None


def check_nextcloud_file_exists(nextcloud_path):
    """Check if a file exists in Nextcloud using WebDAV."""
    try:
        webdav_url = f"{NEXTCLOUD_URL}/remote.php/dav/files/{NEXTCLOUD_USERNAME}/{nextcloud_path}"
        
        response = requests.head(
            webdav_url,
            auth=(NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN),
            timeout=30
        )
        
        # 200 = file exists, 404 = file doesn't exist
        return response.status_code == 200
        
    except Exception as e:
        logger.error(f"Error checking if Nextcloud file exists: {e}")
        return False


def generate_unique_filename(base_filename, folder_path):
    """Generate a unique filename if conflicts exist in Nextcloud."""
    try:
        # Split filename and extension
        name_part, ext = os.path.splitext(base_filename)
        counter = 1
        
        # Check if base filename exists
        test_path = f"{folder_path}/{base_filename}"
        if not check_nextcloud_file_exists(test_path):
            return base_filename
        
        # Try numbered versions until we find one that doesn't exist
        while counter <= 50:  # Reasonable limit to prevent infinite loops
            new_filename = f"{name_part}_v{counter}{ext}"
            test_path = f"{folder_path}/{new_filename}"
            
            if not check_nextcloud_file_exists(test_path):
                logger.info(f"Generated unique filename: {new_filename}")
                return new_filename
            
            counter += 1
        
        # If we still have conflicts after 50 tries, add timestamp
        timestamp = datetime.datetime.now().strftime('%H%M%S')
        unique_filename = f"{name_part}_{timestamp}{ext}"
        logger.warning(f"Used timestamp for unique filename: {unique_filename}")
        return unique_filename
        
    except Exception as e:
        logger.error(f"Error generating unique filename: {e}")
        # Fallback to timestamp-based name
        timestamp = datetime.datetime.now().strftime('%H%M%S')
        name_part, ext = os.path.splitext(base_filename)
        return f"{name_part}_{timestamp}{ext}"


def extract_meeting_datetime(title):
    """Extract meeting date and time from the title text."""
    try:
        # Common date patterns to look for
        date_patterns = [
            # "Jun 25 2025", "June 25, 2025", etc.
            r'([A-Za-z]{3,9})\s+(\d{1,2}),?\s+(\d{4})',
            # "2025-06-25", "06/25/2025", etc.
            r'(\d{4})[-/](\d{1,2})[-/](\d{1,2})',
            r'(\d{1,2})[-/](\d{1,2})[-/](\d{4})'
        ]
        
        # Time patterns to look for
        time_patterns = [
            # "916 AM-1158 AM", "9:16 AM-11:58 AM", etc.
            r'(\d{1,2}):?(\d{2})\s*(AM|PM)[-–]\s*(\d{1,2}):?(\d{2})\s*(AM|PM)',
            # Single time "916 AM", "9:16 AM"
            r'(\d{1,2}):?(\d{2})\s*(AM|PM)'
        ]
        
        meeting_date = None
        meeting_time = None
        
        # Extract date
        for pattern in date_patterns:
            match = re.search(pattern, title, re.IGNORECASE)
            if match:
                groups = match.groups()
                if len(groups) == 3:
                    if groups[0].isdigit():  # Format: YYYY-MM-DD or MM/DD/YYYY
                        if len(groups[0]) == 4:  # YYYY-MM-DD
                            year, month, day = groups
                        else:  # MM/DD/YYYY
                            month, day, year = groups
                    else:  # Format: Month DD YYYY
                        month_name, day, year = groups
                        # Convert month name to number
                        month_mapping = {
                            'jan': '01', 'january': '01',
                            'feb': '02', 'february': '02',
                            'mar': '03', 'march': '03',
                            'apr': '04', 'april': '04',
                            'may': '05', 'may': '05',
                            'jun': '06', 'june': '06',
                            'jul': '07', 'july': '07',
                            'aug': '08', 'august': '08',
                            'sep': '09', 'september': '09',
                            'oct': '10', 'october': '10',
                            'nov': '11', 'november': '11',
                            'dec': '12', 'december': '12'
                        }
                        month = month_mapping.get(month_name.lower()[:3])
                        if not month:
                            continue
                    
                    # Format as MMDDYY
                    try:
                        month_num = int(month)
                        day_num = int(day)
                        year_num = int(year)
                        
                        # Convert to MMDDYY format
                        meeting_date = f"{month_num:02d}{day_num:02d}{year_num % 100:02d}"
                        break
                    except ValueError:
                        continue
        
        # Extract time
        for pattern in time_patterns:
            match = re.search(pattern, title, re.IGNORECASE)
            if match:
                groups = match.groups()
                if len(groups) >= 3:  # At least start time
                    if len(groups) == 6:  # Start and end time
                        start_hour, start_min, start_period, end_hour, end_min, end_period = groups
                        meeting_time = f"{start_hour}{start_min} {start_period.upper()}-{end_hour}{end_min} {end_period.upper()}"
                    else:  # Single time
                        hour, minute, period = groups[:3]
                        meeting_time = f"{hour}{minute} {period.upper()}"
                    break
        
        return meeting_date, meeting_time
        
    except Exception as e:
        logger.warning(f"Error extracting meeting date/time: {e}")
        return None, None


def create_folder_hierarchy(base_folder_path, date_folder):
    """Create the full folder hierarchy including date subfolder."""
    try:
        # Create base folder structure first
        folder_parts = base_folder_path.split('/')
        current_path = ""
        
        for part in folder_parts:
            if current_path:
                current_path += f"/{part}"
            else:
                current_path = part
            
            success = create_nextcloud_folder(current_path)
            if not success:
                logger.warning(f"Could not create folder: {current_path}")
        
        # Create date subfolder
        full_path = f"{base_folder_path}/{date_folder}"
        success = create_nextcloud_folder(full_path)
        
        return success, full_path
        
    except Exception as e:
        logger.error(f"Error creating folder hierarchy: {e}")
        return False, None


def _add_formatted_transcript_to_doc(doc, transcript_text):
    """
    Parse transcript text and add to document with bold speaker labels.
    Handles format like "Speaker A: text" with proper bold formatting.
    """
    import re
    
    # Split transcript into paragraphs (double line breaks separate speakers)
    paragraphs = transcript_text.split('\n\n')
    
    for paragraph_text in paragraphs:
        paragraph_text = paragraph_text.strip()
        if not paragraph_text:
            continue
            
        # Check if paragraph starts with speaker label pattern
        speaker_match = re.match(r'^((?:\[[^\]]+\]\s+)?Speaker [A-Z]):(.*)$', paragraph_text, re.DOTALL)
        
        if speaker_match:
            speaker_label = speaker_match.group(1).strip()
            text_content = speaker_match.group(2).strip()
            
            # Create paragraph
            para = doc.add_paragraph()
            
            # Add speaker label in bold
            speaker_run = para.add_run(f"{speaker_label}: ")
            speaker_run.bold = True
            
            # Add the text content
            para.add_run(text_content)
        else:
            # No speaker label found, add as regular paragraph
            doc.add_paragraph(paragraph_text)


def create_docx_document_with_metadata(transcript_data, title, entry_url, meeting_info):
    """Create a DOCX document with enhanced meeting metadata."""
    try:
        # Create a new document
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add meeting information section
        doc.add_heading('Meeting Information', level=1)
        
        # Committee information
        if meeting_info['type'] == 'interim':
            info_paragraph = doc.add_paragraph()
            info_paragraph.add_run('Meeting Type: ').bold = True
            info_paragraph.add_run('Interim Committee')
            
        elif meeting_info['type'] == 'session':
            info_paragraph = doc.add_paragraph()
            info_paragraph.add_run('Meeting Type: ').bold = True
            info_paragraph.add_run('Legislative Session Committee')
            
            info_paragraph = doc.add_paragraph()
            info_paragraph.add_run('Chamber: ').bold = True
            info_paragraph.add_run(meeting_info['chamber'])
        
        info_paragraph = doc.add_paragraph()
        info_paragraph.add_run('Committee: ').bold = True
        info_paragraph.add_run(meeting_info['committee_name'])
        
        # Standard metadata
        info_paragraph = doc.add_paragraph()
        info_paragraph.add_run('Source URL: ').bold = True
        info_paragraph.add_run(entry_url)
        
        info_paragraph = doc.add_paragraph()
        info_paragraph.add_run('Transcribed by: ').bold = True
        info_paragraph.add_run('LWE.Vote')
        
        info_paragraph = doc.add_paragraph()
        info_paragraph.add_run('Transcribed on: ').bold = True
        info_paragraph.add_run(datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        
        if transcript_data and isinstance(transcript_data, dict):
            info_paragraph = doc.add_paragraph()
            info_paragraph.add_run('Audio duration: ').bold = True
            info_paragraph.add_run(f"{transcript_data.get('audio_duration', 'Unknown')} seconds")
            
            info_paragraph = doc.add_paragraph()
            info_paragraph.add_run('Confidence score: ').bold = True
            info_paragraph.add_run(f"{transcript_data.get('confidence', 'Unknown')}")
        
        # Add separator
        doc.add_paragraph('_' * 80)
        
        # Add transcript section
        doc.add_heading('Transcript', level=1)
        
        # Format and add transcript content
        if isinstance(transcript_data, str):
            # Handle formatted transcript string - parse and format speaker labels
            _add_formatted_transcript_to_doc(doc, transcript_data)
        elif transcript_data and isinstance(transcript_data, dict):
            # Handle structured transcript data
            if "words" in transcript_data:
                current_speaker = None
                current_paragraph = None
                word_buffer = []
                
                for i, word in enumerate(transcript_data["words"]):
                    speaker = word.get("speaker", "Unknown Speaker")
                    word_text = word.get("text", "")
                    
                    # Check if speaker changed
                    if speaker != current_speaker:
                        # Finish previous paragraph if exists
                        if word_buffer and current_paragraph:
                            current_paragraph.add_run(' '.join(word_buffer) + ' ')
                            word_buffer = []
                        
                        # Start new paragraph for new speaker
                        current_paragraph = doc.add_paragraph()
                        current_speaker = speaker
                        
                        # Add speaker name in bold
                        speaker_run = current_paragraph.add_run(f"{speaker}: ")
                        speaker_run.bold = True
                        
                        # Start collecting words for this speaker
                        word_buffer = [word_text]
                    else:
                        # Same speaker, continue collecting words
                        word_buffer.append(word_text)
                    
                    # Create natural sentence breaks
                    if (word_text.endswith(('.', '!', '?')) and len(word_buffer) > 5) or len(word_buffer) >= 50:
                        # Add current buffer to paragraph and start fresh
                        if current_paragraph:
                            current_paragraph.add_run(' '.join(word_buffer) + ' ')
                            word_buffer = []
                    
                    # Prevent overly long paragraphs - create new paragraph for same speaker
                    if len(' '.join(word_buffer)) > 500:  # About 500 characters
                        if current_paragraph:
                            current_paragraph.add_run(' '.join(word_buffer))
                            word_buffer = []
                            
                            # Start new paragraph for same speaker (continuation)
                            current_paragraph = doc.add_paragraph()
                            speaker_run = current_paragraph.add_run(f"{current_speaker} (continued): ")
                            speaker_run.bold = True
                
                # Add any remaining words
                if word_buffer and current_paragraph:
                    current_paragraph.add_run(' '.join(word_buffer))
            elif "text" in transcript_data:
                # Fallback to plain text if words are not available
                doc.add_paragraph(transcript_data["text"])
            else:
                doc.add_paragraph("No transcription text available.")
        else:
            doc.add_paragraph("No transcription text available.")
        
        return doc
        
    except Exception as e:
        logger.error(f"Error creating DOCX document with metadata: {e}")
        return None


def save_transcript_to_nextcloud(transcript_data, title, entry_url, save_as_txt=False):
    """Save transcript to Nextcloud with structured folder organization.

    Args:
        transcript_data: Transcript content (dict for structured data, string for plain text)
        title: Meeting title
        entry_url: Source URL
        save_as_txt: If True, save as plain .txt file. If False, create .docx document (default)
    """
    try:
        if not all([NEXTCLOUD_URL, NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN]):
            logger.error("Nextcloud configuration missing")
            return None
        
        # Parse the meeting title to determine folder structure
        meeting_info = parse_meeting_title(title)
        
        # Extract meeting date and time from title
        meeting_date, meeting_time = extract_meeting_datetime(title)
        
        # Use extracted date for folder, fallback to current date
        if meeting_date:
            # Convert MMDDYY to YYYY-MM-DD for folder structure
            mm = meeting_date[:2]
            dd = meeting_date[2:4]
            yy = meeting_date[4:6]
            # Assume 20xx for years
            yyyy = f"20{yy}" if int(yy) < 50 else f"19{yy}"
            date_folder = f"{yyyy}-{mm}-{dd}"
            logger.info(f"Using extracted meeting date: {date_folder}")
        else:
            date_folder = datetime.datetime.now().strftime('%Y-%m-%d')
            logger.warning(f"Could not extract meeting date from title, using current date: {date_folder}")
        
        # Create the full folder hierarchy
        success, full_folder_path = create_folder_hierarchy(meeting_info['folder_path'], date_folder)
        
        if not success:
            logger.error("Failed to create folder hierarchy, falling back to simple structure")
            # Fallback to simple structure
            simple_folder = f"Legislative Transcription/{date_folder}"
            create_nextcloud_folder("Legislative Transcription")
            create_nextcloud_folder(simple_folder)
            full_folder_path = simple_folder
        
        # Create base filename using the new format: YYYYMMDD-IC-ACRONYM-TIME.docx
        if meeting_date:
            # Convert MMDDYY to YYYYMMDD
            mm = meeting_date[:2]
            dd = meeting_date[2:4]
            yy = meeting_date[4:6]
            yyyy = f"20{yy}" if int(yy) < 50 else f"19{yy}"
            date_prefix = f"{yyyy}{mm}{dd}"
        else:
            # Fallback to current date in YYYYMMDD format
            now = datetime.datetime.now()
            date_prefix = f"{now.year:04d}{now.month:02d}{now.day:02d}"
        
        # Get meeting type prefix and committee acronym
        meeting_prefix = meeting_info.get('prefix', 'UNKNOWN')  # IC, HOUSE, SENATE, etc.
        committee_acronym = meeting_info.get('committee_acronym', 'UNK')
        
        # Build filename: YYYYMMDD-PREFIX-ACRONYM-TIME.docx
        filename_parts = [date_prefix, meeting_prefix, committee_acronym]
        
        if meeting_time:
            # Clean up time format - remove spaces and standardize
            # Convert "904AM-133PM" format to "904AM-133PM" (already clean)
            clean_time = meeting_time.replace(' ', '').replace('–', '-')  # Handle different dash types
            filename_parts.append(clean_time)
            logger.info(f"Using extracted meeting time: {clean_time}")
        else:
            logger.warning("Could not extract meeting time from title")
        
        # Join parts with dashes and clean for filename
        filename_base = "-".join(filename_parts)
        # Clean for filesystem (remove problematic characters)
        filename_base = re.sub(r'[<>:"/\\|?*]', '', filename_base)

        # Choose file extension based on format
        file_extension = '.txt' if save_as_txt else '.docx'
        base_filename = f"{filename_base}{file_extension}"
        
        logger.info(f"Generated filename: {base_filename}")
        
        # Generate unique filename if conflicts exist
        unique_filename = generate_unique_filename(base_filename, full_folder_path)

        # Save to temporary file
        temp_path = os.path.join(DOWNLOAD_DIR, unique_filename)

        if save_as_txt:
            # Save as plain text file
            with open(temp_path, 'w', encoding='utf-8') as f:
                if isinstance(transcript_data, str):
                    f.write(transcript_data)
                elif isinstance(transcript_data, dict):
                    # Extract text from dict if needed
                    f.write(transcript_data.get('text', str(transcript_data)))
                else:
                    f.write(str(transcript_data))
            logger.info(f"Created plain text file: {temp_path}")
        else:
            # Create the DOCX document with enhanced metadata
            doc = create_docx_document_with_metadata(transcript_data, title, entry_url, meeting_info)
            if not doc:
                logger.error("Failed to create DOCX document")
                return None
            doc.save(temp_path)
            logger.info(f"Created DOCX document: {temp_path}")
        
        # Upload to Nextcloud
        nextcloud_file_path = f"{full_folder_path}/{unique_filename}"
        
        if upload_to_nextcloud(temp_path, nextcloud_file_path):
            logger.info(f"Successfully saved transcript to: {nextcloud_file_path}")

            # Send webhook notification to n8n
            send_n8n_webhook(unique_filename, full_folder_path)

            # Create share link
            share_link = get_nextcloud_share_link(nextcloud_file_path)
            
            # Clean up temporary file
            try:
                os.remove(temp_path)
            except Exception as e:
                logger.warning(f"Could not remove temporary file: {e}")
            
            return {
                'nextcloud_path': nextcloud_file_path,
                'share_link': share_link,
                'filename': unique_filename,
                'meeting_info': meeting_info,
                'folder_path': full_folder_path,
                'meeting_date': meeting_date,
                'meeting_time': meeting_time
            }
        else:
            logger.error("Failed to upload to Nextcloud")
            return None
            
    except Exception as e:
        logger.error(f"Error saving transcript to Nextcloud: {e}")
        return None
