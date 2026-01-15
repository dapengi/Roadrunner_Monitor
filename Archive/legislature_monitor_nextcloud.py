#!/usr/bin/env python3.12
"""
New Mexico Legislature Monitoring Script with yt-dlp, Oxylabs Proxy, and Nextcloud Integration
Updated to use Nextcloud instead of Google Drive/Docs
"""

import requests
from bs4 import BeautifulSoup
import time
import schedule
import os
import datetime
import smtplib
import json
import re
import urllib.parse
import subprocess
import glob
import random
import yt_dlp
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from dotenv import load_dotenv
import logging
from pathlib import Path
import tempfile
from io import BytesIO

# Document creation imports
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("legislature_monitor.log"),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger(__name__)

# Load environment variables from .env file
load_dotenv()

# URL to monitor
URL = "https://sg001-harmony.sliq.net/00293/Harmony/en/View/RecentEnded/20250522/-1"

# File to store the latest entries
ENTRIES_FILE = "latest_entries.txt"
PROCESSED_ENTRIES_FILE = "processed_entries.txt"
DOWNLOAD_DIR = "downloads"
LAST_CLEANUP_FILE = "last_cleanup.txt"
PROXY_LIST_FILE = "proxy_list.txt"
LAST_PROXY_UPDATE_FILE = "last_proxy_update.txt"

# Email configuration
EMAIL_USER = os.getenv("EMAIL_USER")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")
EMAIL_RECIPIENTS = ["joshua@lwe.digital", "skyedevore@gmail.com"]

# AssemblyAI API configuration
ASSEMBLYAI_API_KEY = os.getenv("ASSEMBLYAI_API_KEY")
ASSEMBLYAI_BASE_URL = "https://api.assemblyai.com/v2"

# Oxylabs Proxy configuration
OXYLABS_PROXY_HOST = "dc.oxylabs.io"
OXYLABS_PROXY_PORT = "8000"
OXYLABS_USERNAME = os.getenv("OXYLABS_USERNAME", "user-dapengi_IAs3C-country-US")
OXYLABS_PASSWORD = os.getenv("OXYLABS_PASSWORD", "XDVLR8SYrj2Ee7+")
OXYLABS_LOCATION_URL = "https://ip.oxylabs.io/location"
PROXY_UPDATE_INTERVAL_HOURS = 6  # Check proxy status every 6 hours

# Nextcloud configuration
NEXTCLOUD_URL = os.getenv("NEXTCLOUD_URL")
NEXTCLOUD_USERNAME = os.getenv("NEXTCLOUD_USERNAME")
NEXTCLOUD_TOKEN = os.getenv("NEXTCLOUD_TOKEN")
NEXTCLOUD_BASE_FOLDER = "Legislative Transcription"

# Create download directory if it doesn't exist
Path(DOWNLOAD_DIR).mkdir(exist_ok=True)

# Global proxy configuration
proxy_config = None
proxy_working = False


def test_oxylabs_proxy():
    """Test if Oxylabs proxy is working and get IP location info."""
    global proxy_config, proxy_working
    
    try:
        logger.info("Testing Oxylabs proxy connection...")
        
        # Create proxy configuration using working method (Method 2 from test)
        proxy_config = {
            'ip': OXYLABS_PROXY_HOST,
            'port': OXYLABS_PROXY_PORT,
            'username': OXYLABS_USERNAME,
            'password': OXYLABS_PASSWORD
        }
        
        # Test the proxy using the working format from our test
        proxy_url = f"http://{OXYLABS_USERNAME}:{OXYLABS_PASSWORD}@{OXYLABS_PROXY_HOST}:{OXYLABS_PROXY_PORT}"
        proxies = {
            "http": proxy_url,
            "https": proxy_url
        }
        
        response = requests.get(
            OXYLABS_LOCATION_URL,
            proxies=proxies,
            timeout=30
        )
        response.raise_for_status()
        
        location_data = response.json()
        
        # Parse the complex Oxylabs response format
        proxy_ip = location_data.get('ip', 'unknown')
        
        # Extract country/city from the providers (try multiple sources)
        country = 'unknown'
        city = 'unknown'
        
        providers = location_data.get('providers', {})
        for provider_name, provider_data in providers.items():
            if provider_data.get('country') and country == 'unknown':
                country = provider_data.get('country', 'unknown')
            if provider_data.get('city') and city == 'unknown':
                city = provider_data.get('city', 'unknown')
        
        logger.info(f"Oxylabs proxy working - IP: {proxy_ip}, Country: {country}, City: {city}")
        
        proxy_working = True
        
        # Save proxy test result
        with open(LAST_PROXY_UPDATE_FILE, 'w') as f:
            f.write(datetime.datetime.now().isoformat())
        
        return True
        
    except Exception as e:
        logger.error(f"Oxylabs proxy test failed: {e}")
        proxy_working = False
        return False


def should_update_proxy_list():
    """Check if we should test the proxy again."""
    try:
        if not os.path.exists(LAST_PROXY_UPDATE_FILE):
            return True
        
        with open(LAST_PROXY_UPDATE_FILE, 'r') as f:
            last_update_str = f.read().strip()
        
        last_update = datetime.datetime.fromisoformat(last_update_str)
        now = datetime.datetime.now()
        
        hours_since_update = (now - last_update).total_seconds() / 3600
        return hours_since_update >= PROXY_UPDATE_INTERVAL_HOURS
        
    except Exception as e:
        logger.error(f"Error checking proxy update time: {e}")
        return True


def get_proxy_dict(proxy):
    """Convert proxy info to requests-compatible proxy dict using working Oxylabs format."""
    if not proxy:
        return None
    
    if proxy['username'] and proxy['password']:
        # Working format: http://username:password@host:port (Method 2 from test)
        proxy_url = f"http://{proxy['username']}:{proxy['password']}@{proxy['ip']}:{proxy['port']}"
    else:
        # Without authentication
        proxy_url = f"http://{proxy['ip']}:{proxy['port']}"
    
    return {
        'http': proxy_url,
        'https': proxy_url
    }


def get_proxy_url(proxy):
    """Convert proxy info to URL format for yt-dlp using working Oxylabs format."""
    if not proxy:
        return None
    
    if proxy['username'] and proxy['password']:
        return f"http://{proxy['username']}:{proxy['password']}@{proxy['ip']}:{proxy['port']}"
    else:
        return f"http://{proxy['ip']}:{proxy['port']}"


def make_request_with_proxy(url, headers=None, max_retries=3):
    """Make a request using Oxylabs proxy with retry logic."""
    if headers is None:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
                          'AppleWebKit/537.36 (KHTML, like Gecko) '
                          'Chrome/91.0.4472.124 Safari/537.36'
        }
    
    for attempt in range(max_retries):
        try:
            if proxy_working and proxy_config:
                proxy_dict = get_proxy_dict(proxy_config)
                logger.info(f"Attempt {attempt + 1}: Using Oxylabs proxy {proxy_config['ip']}:{proxy_config['port']}")
                response = requests.get(url, headers=headers, proxies=proxy_dict, timeout=30)
            else:
                logger.warning(f"Attempt {attempt + 1}: Proxy not available, using direct connection")
                response = requests.get(url, headers=headers, timeout=30)
            
            response.raise_for_status()
            logger.info(f"Request successful on attempt {attempt + 1}")
            return response
            
        except requests.exceptions.RequestException as e:
            logger.warning(f"Request failed on attempt {attempt + 1}: {e}")
            if attempt < max_retries - 1:
                # Wait before retry with exponential backoff
                wait_time = (2 ** attempt) + random.uniform(1, 3)
                logger.info(f"Waiting {wait_time:.1f} seconds before retry...")
                time.sleep(wait_time)
            else:
                logger.error(f"All {max_retries} attempts failed for URL: {url}")
                raise
    
    return None


def cleanup_downloads_directory():
    """Delete all files in the downloads directory."""
    try:
        files_pattern = os.path.join(DOWNLOAD_DIR, "*")
        files_to_delete = glob.glob(files_pattern)
        
        if not files_to_delete:
            logger.info("Downloads directory is already empty")
            return
        
        deleted_count = 0
        total_size = 0
        
        for file_path in files_to_delete:
            try:
                if os.path.isfile(file_path):
                    file_size = os.path.getsize(file_path)
                    total_size += file_size
                    os.remove(file_path)
                    deleted_count += 1
                    logger.info(f"Deleted: {os.path.basename(file_path)} ({file_size / (1024*1024):.2f} MB)")
                    
            except Exception as e:
                logger.error(f"Error deleting {file_path}: {e}")
        
        logger.info(f"Cleanup completed: {deleted_count} files deleted, {total_size / (1024*1024):.2f} MB freed")
        
        with open(LAST_CLEANUP_FILE, 'w') as f:
            f.write(datetime.datetime.now().strftime('%Y-%m-%d'))
        
    except Exception as e:
        logger.error(f"Error during downloads cleanup: {e}")


def should_run_daily_cleanup():
    """Check if daily cleanup should be run."""
    try:
        today = datetime.datetime.now().strftime('%Y-%m-%d')
        
        if not os.path.exists(LAST_CLEANUP_FILE):
            return True
        
        with open(LAST_CLEANUP_FILE, 'r') as f:
            last_cleanup_date = f.read().strip()
        
        return last_cleanup_date != today
        
    except Exception as e:
        logger.error(f"Error checking cleanup date: {e}")
        return True


def get_committee_rosters():
    """Return comprehensive committee member rosters with nickname handling."""
    return {
        'Water & Natural Resources': [
            'Liz Stefanics', 'Matthew McQueen', 'Micaela Lara Cadena', 'Pete Campos',
            'Kathleen Cates', 'Joseph Cervantes', 'Candy Spence Ezzell', 'Miguel García',
            'Jonathan Henry', 'Joseph Franklin Hernandez', 'Wonda Johnson', 'Marian Matthews',
            'Debbie O\'Malley', 'Kristina Ortez', 'Angelica Rubio', 'Debra Sariñana',
            'Larry Scott', 'Elaine Sena Cortez', 'Jeff Steinborn', 'Jim Townsend',
            'Martin Zamora', 'Gail Armstrong', 'Craig Brandt', 'Angel Charley',
            'Jack Chatfield', 'Nicole Chavez', 'Eleanor Chávez', 'Meredith Dixon',
            'Joanne Ferrary', 'Martha Garcia', 'Charlotte Little', 'Linda López',
            'Stefani Lord', 'Tara Lujan', 'Moe Maestas', 'Javier Martínez',
            'Angelita Mejia', 'Tanya Mirabal Moya', 'Rod Montoya', 'Mark Murphy',
            'Randall Pettigrew', 'Patricia Roybal Caballero', 'Antoinette Sedillo Lopez',
            'Nathan Small', 'Mimi Stewart', 'Anthony Thornton', 'Harlan Vincent',
            'Peter Wirth', 'Pat Woods'
        ],
        'Transportation Infrastructure Revenue': [
            'Art De La Cruz', 'Moe Maestas', 'Cynthia Borrego', 'Craig Brandt',
            'Cathrynn Brown', 'David Gallegos', 'Bobby Gonzales', 'Day Hochman-Vigil',
            'Jenifer Jones', 'Ray Lara', 'Patty Lundstrom', 'George Muñoz',
            'Debbie O\'Malley', 'Randy Pettigrew', 'Gail Armstrong', 'Heather Berghmans',
            'Pat Boone', 'Rebecca Dow', 'Martha Garcia', 'Joy Garratt',
            'Anita Gonzales', 'Rod Montoya', 'Sarah Silva', 'Martin Zamora'
        ],
        'Tobacco Settlement Revenue Oversight': [
            'Joanne Ferrary', 'Martin Hickey', 'Brian Baca', 'Linda López',
            'Gabriel Ramos', 'Liz Thomson', 'Shannon Pinto', 'Larry Scott', 'Luis Terrazas'
        ],
        'Science, Technology & Telecommunications': [
            'Debra Sariñana', 'Heather Berghmans', 'Joy Garratt', 'Anita Gonzales',
            'Joshua Hernandez', 'Tara Lujan', 'Alan Martinez', 'Michael Padilla',
            'Harold Pope', 'Bill Soules', 'Anthony Thornton', 'Nicole Tobiassen',
            'Marianna Anaya', 'Gail Armstrong', 'Jay Block', 'Christine Chandler',
            'Meredith Dixon', 'Yanira Gurrola', 'Charlotte Little', 'Larry Scott',
            'Sarah Silva', 'Jim Townsend', 'Linda Trujillo'
        ],
        'Revenue Stabilization & Tax Policy': [
            'Carrie Hamblen', 'Derrick Lente', 'Heather Berghmans', 'Craig Brandt',
            'Micaela Lara Cadena', 'Pete Campos', 'Christine Chandler', 'Mark Duncan',
            'Natalie Figueroa', 'Doreen Gallegos', 'Patty Lundstrom', 'George Muñoz',
            'Mark Murphy', 'Cristina Parajón', 'Gabriel Ramos', 'Bill Sharer',
            'Luis Terrazas', 'Peter Wirth', 'Joanne Ferrary', 'Moe Maestas',
            'Alan Martinez', 'Kristina Ortez', 'Antoinette Sedillo Lopez',
            'Elaine Sena Cortez', 'Jim Townsend', 'Jonathan Henry', 'Joshua Hernandez',
            'Leo Jaramillo', 'Javier Martínez', 'Rod Montoya', 'Debbie O\'Malley',
            'Nicholas Paul', 'Joshua Sanchez', 'Sarah Silva'
        ],
        'Radioactive & Hazardous Materials': [
            'Jeff Steinborn', 'Joanne Ferrary', 'Cathrynn Brown', 'Angel Charley',
            'Meredith Dixon', 'David Gallegos', 'Yanira Gurrola', 'Wonda Johnson',
            'Stefani Lord', 'Harold Pope', 'Antoinette Sedillo Lopez', 'Jim Townsend',
            'John Block', 'Crystal Brantley', 'Shannon Pinto', 'Debra Sariñana',
            'Anthony Thornton'
        ],
        'Public School Capital Outlay Oversight': [
            'Andrés Romero', 'Bill Soules', 'Pat Boone', 'Catherine Cullen',
            'Steve Lanier', 'Derrick Lente', 'Javier Martínez', 'Tanya Mirabal Moya',
            'George Muñoz', 'Shannon Pinto', 'Debra Sariñana', 'Nathan Small',
            'Mimi Stewart', 'Julie Lucero', 'Cindy Montoya', 'Brandy Murphy',
            'Mariana Padilla', 'Wayne Propst', 'Amber Romero', 'Stan Rounds',
            'Ignacio Ruiz', 'Rhiannon Chavez', 'Steve Carlson', 'Johnny Benavidez',
            'LeAnne Gandy', 'Ashley Leach', 'Antonio Ortiz'
        ],
        'New Mexico Finance Authority Oversight': [
            'Joy Garratt', 'Michael Padilla', 'Michelle Paulene Abeyta', 'John Block',
            'Cynthia Borrego', 'Craig Brandt', 'Kathleen Cates', 'Jack Chatfield',
            'Meredith Dixon', 'Wonda Johnson', 'Tara Lujan', 'Angelita Mejia',
            'Debbie O\'Malley', 'Kristina Ortez', 'Gabriel Ramos',
            'Patricia Roybal Caballero', 'Luis Terrazas', 'Susan Herrera',
            'Steve Lanier', 'Stefani Lord', 'Debra Sariñana', 'Sarah Silva',
            'Nicole Tobiassen'
        ],
        'Mortgage Finance Authority Act Oversight': [
            'Janelle Anyanonu', 'Linda Trujillo', 'Craig Brandt', 'Bobby Gonzales',
            'Marian Matthews', 'Rod Montoya', 'Sarah Silva', 'Nicole Tobiassen',
            'Gail Armstrong', 'Cynthia Borrego', 'Kathleen Cates', 'Eleanor Chávez',
            'Meredith Dixon', 'Rebecca Dow', 'Jonathan Henry', 'Wonda Johnson',
            'Tara Lujan', 'Cindy Nava', 'Debbie O\'Malley', 'Michael Padilla',
            'Harold Pope', 'Patricia Roybal Caballero', 'Luis Terrazas'
        ],
        'Military & Veterans\' Affairs': [
            'Harold Pope', 'Debra Sariñana', 'Jay Block', 'Art De La Cruz',
            'Wonda Johnson', 'Nicholas Paul', 'Shannon Pinto', 'Luis Terrazas',
            'Pat Boone', 'Craig Brandt', 'Catherine Cullen', 'Stefani Lord',
            'Alan Martinez', 'Angelita Mejia', 'Gabriel Ramos', 'Andrea Reeb'
        ],
        'Legislative Health & Human Services': [
            'Liz Thomson', 'Linda López', 'Jay Block', 'Eleanor Chávez',
            'Pamelya Herndon', 'Martin Hickey', 'Jenifer Jones', 'Larry Scott',
            'Marianna Anaya', 'John Block', 'Kathleen Cates', 'Christine Chandler',
            'Angel Charley', 'Nicole Chavez', 'Katy Duhigg', 'Joanne Ferrary',
            'David Gallegos', 'Miguel García', 'Anita Gonzales', 'Wonda Johnson',
            'Tara Lujan', 'Alan Martinez', 'Shannon Pinto', 'Harold Pope',
            'Patricia Roybal Caballero', 'Antoinette Sedillo Lopez', 'Elaine Sena Cortez',
            'Sarah Silva', 'Liz Stefanics', 'Reena Szczepanski', 'Diane Torres-Velásquez'
        ],
        'Legislative Finance': [
            'Nathan Small', 'George Muñoz', 'Pete Campos', 'Jack Chatfield',
            'Meredith Dixon', 'Rebecca Dow', 'Bobby Gonzales', 'Susan Herrera',
            'Steve Lanier', 'Derrick Lente', 'Joseph Sanchez', 'Benny Shendo',
            'Nicole Tobiassen', 'Linda Trujillo', 'Harlan Vincent', 'Pat Woods',
            'Gail Armstrong', 'Eleanor Chávez', 'Art De La Cruz', 'David Gallegos',
            'Jonathan Henry', 'Martin Hickey', 'Day Hochman-Vigil', 'Kristina Ortez',
            'Andrea Reeb', 'Bill Sharer', 'Bill Soules', 'Liz Stefanics',
            'Reena Szczepanski', 'Luis Terrazas', 'Peter Wirth', 'Brian Baca',
            'Craig Brandt', 'Cathrynn Brown', 'Mark Duncan', 'Joy Garratt',
            'Joseph Hernandez', 'Pamelya Herndon', 'Charlotte Little', 'Tara Lujan',
            'Michael Padilla', 'Randall Pettigrew', 'Debra Sariñana', 'Sarah Silva',
            'Jeff Steinborn'
        ],
        'Legislative Education Study': [
            'Bill Soules', 'Andrés Romero', 'Brian Baca', 'Craig Brandt',
            'Candy Spence Ezzell', 'Joy Garratt', 'Ray Lara', 'Tanya Mirabal Moya',
            'Debra Sariñana', 'Mimi Stewart', 'Gail Armstrong', 'John Block',
            'Pat Boone', 'Crystal Brantley', 'Angel Charley', 'Jack Chatfield',
            'Catherine Cullen', 'Natalie Figueroa', 'Yanira Gurrola', 'Jonathan Henry',
            'Wonda Johnson', 'Linda López', 'Cindy Nava', 'Harold Pope',
            'Gabriel Ramos', 'Patricia Roybal Caballero', 'Nathan Small',
            'Anthony Thornton', 'Diane Torres-Velásquez', 'Harlan Vincent'
        ],
        'Legislative Council': [
            'Javier Martínez', 'Mimi Stewart', 'Gail Armstrong', 'Pete Campos',
            'Christine Chandler', 'Rebecca Dow', 'Katy Duhigg', 'David Gallegos',
            'Doreen Gallegos', 'Day Hochman-Vigil', 'Alan Martinez', 'Michael Padilla',
            'Bill Sharer', 'Reena Szczepanski', 'Peter Wirth', 'Pat Woods',
            'Craig Brandt', 'Angel Charley', 'Natalie Figueroa', 'Wonda Johnson',
            'Patricia Roybal Caballero', 'Liz Thomson'
        ],
        'Land Grant': [
            'Linda Serrato', 'Leo Jaramillo', 'Angel Charley', 'David Gallegos',
            'Miguel García', 'Susan Herrera', 'Stefani Lord', 'Moe Maestas',
            'Matthew McQueen', 'Angelita Mejia', 'Liz Stefanics', 'Anthony Thornton',
            'Diane Torres-Velásquez', 'Catherine Cullen', 'Anita Gonzales',
            'Wonda Johnson', 'Linda López', 'Jimmy Mason', 'Rod Montoya',
            'Debbie O\'Malley', 'Joshua Sanchez'
        ],
        'Investments & Pensions Oversight': [
            'Cynthia Borrego', 'Bobby Gonzales', 'Mark Duncan', 'Natalie Figueroa',
            'Bill Hall', 'Tara Lujan', 'Antoinette Sedillo Lopez', 'Linda Serrato',
            'Sarah Silva', 'Liz Stefanics', 'Pat Woods', 'Gail Armstrong',
            'Stefani Lord', 'George Muñoz', 'Patricia Roybal Caballero', 'Mimi Stewart',
            'Anthony Thornton'
        ],
        'Interim Legislative Ethics': [
            'Andrea Romero', 'Liz Stefanics', 'Cathrynn Brown', 'Christine Chandler',
            'Bobby Gonzales', 'Bill Hall', 'Carrie Hamblen', 'Leo Jaramillo',
            'Patty Lundstrom', 'Nicholas Paul', 'Andrea Reeb', 'Patricia Roybal Caballero',
            'Bill Sharer', 'Jim Townsend', 'Pat Woods', 'Kathleen Cates',
            'Catherine Cullen', 'Yanira Gurrola'
        ],
        'Indian Affairs': [
            'Shannon Pinto', 'Wonda Johnson', 'Michelle Paulene Abeyta', 'John Block',
            'Angel Charley', 'Martha Garcia', 'Joseph Franklin Hernandez',
            'Charlotte Little', 'Cindy Nava', 'Nicholas Paul', 'Patricia Roybal Caballero',
            'Bill Sharer', 'Benny Shendo', 'Martin Zamora', 'Rebecca Dow',
            'Yanira Gurrola', 'Linda López', 'Stefani Lord', 'Patty Lundstrom',
            'Angelita Mejia', 'George Muñoz', 'Joshua Sanchez', 'Debra Sariñana',
            'Bill Soules', 'Liz Thomson'
        ],
        'Federal Funding Stabilization': [
            'Patty Lundstrom', 'Bill Soules', 'Cathrynn Brown', 'Mark Duncan',
            'Susan Herrera', 'Ray Lara', 'George Muñoz', 'Jim Townsend',
            'Linda Trujillo', 'Pat Woods', 'Meredith Dixon', 'Linda López',
            'Cindy Nava', 'Shannon Pinto', 'Reena Szczepanski'
        ],
        'Economic & Rural Development & Policy': [
            'Angel Charley', 'Doreen Gallegos', 'Janelle Anyanonu', 'Gail Armstrong',
            'Jay Block', 'Pat Boone', 'Crystal Brantley', 'Nicole Chavez',
            'Martha Garcia', 'Joshua Hernandez', 'Moe Maestas', 'Alan Martinez',
            'Marian Matthews', 'Debbie O\'Malley', 'Shannon Pinto',
            'Patricia Roybal Caballero', 'Linda Serrato', 'Sarah Silva', 'Liz Stefanics',
            'Martin Zamora', 'Heather Berghmans', 'Catherine Cullen', 'Meredith Dixon',
            'Natalie Figueroa', 'David Gallegos', 'Anita Gonzales', 'Jenifer Jones',
            'Steve Lanier', 'Ray Lara', 'Charlotte Little', 'Stefani Lord',
            'Tara Lujan', 'Patty Lundstrom', 'Jimmy Mason', 'Rod Montoya',
            'Mark Murphy', 'Joshua Sanchez', 'Bill Sharer', 'Nicole Tobiassen',
            'Jim Townsend'
        ],
        'Courts, Corrections & Justice': [
            'Joseph Cervantes', 'Christine Chandler', 'Marianna Anaya', 'Janelle Anyanonu',
            'Crystal Brantley', 'Nicole Chavez', 'Bill Hall', 'Day Hochman-Vigil',
            'Moe Maestas', 'Cindy Nava', 'Nicholas Paul', 'Andrea Reeb',
            'Andrea Romero', 'Angelica Rubio', 'Antoinette Sedillo Lopez',
            'Gail Armstrong', 'Heather Berghmans', 'John Block', 'Cynthia Borrego',
            'Eleanor Chávez', 'Katy Duhigg', 'Joanne Ferrary', 'Natalie Figueroa',
            'Joy Garratt', 'Yanira Gurrola', 'Steve Lanier', 'Stefani Lord',
            'Tara Lujan', 'Alan Martinez', 'Debbie O\'Malley', 'Debra Sariñana',
            'Sarah Silva', 'Nathan Small', 'Mimi Stewart', 'Reena Szczepanski',
            'Anthony Thornton', 'Linda Trujillo', 'Peter Wirth'
        ],
        'Capitol Security': [
            'Gail Armstrong', 'Alan Martinez', 'Javier Martínez', 'Bill Sharer',
            'Mimi Stewart', 'Reena Szczepanski', 'Peter Wirth', 'Pat Woods'
        ],
        'Capitol Buildings Planning': [
            'Gail Armstrong', 'Javier Martínez', 'Bill Sharer', 'Mimi Stewart',
            'Stephanie Garcia Richard', 'Debra Garcia y Griego', 'Laura Montoya',
            'Wayne Propst', 'Ricky Serna', 'David Thomson'
        ]
    }


def normalize_name(name):
    """Normalize a name for matching (remove quotes, standardize spacing)."""
    # Remove quotes around nicknames and clean up
    name = re.sub(r'["\']', '', name)
    # Standardize spacing
    name = re.sub(r'\s+', ' ', name).strip()
    return name


def get_committee_members_for_meeting(meeting_title):
    """Get the member list for a specific committee meeting."""
    try:
        meeting_info = parse_meeting_title(meeting_title)
        committee_name = meeting_info.get('committee_name', '')
        
        rosters = get_committee_rosters()
        
        # Try exact match first
        if committee_name in rosters:
            return rosters[committee_name]
        
        # Try partial matching for committee names
        for roster_name, members in rosters.items():
            if committee_name.lower() in roster_name.lower() or roster_name.lower() in committee_name.lower():
                logger.info(f"Matched committee '{committee_name}' to roster '{roster_name}'")
                return members
        
        logger.warning(f"No committee roster found for: {committee_name}")
        return []
        
    except Exception as e:
        logger.error(f"Error getting committee members: {e}")
        return []


def identify_speakers_in_transcript(transcript_data, committee_members, confidence_threshold=70):
    """Identify speakers in transcript using committee roster and confidence scoring."""
    try:
        if not transcript_data or 'words' not in transcript_data:
            return transcript_data
        
        # Group words by speaker to analyze speaking patterns
        speakers = {}
        for word in transcript_data['words']:
            speaker = word.get('speaker')
            if speaker:
                if speaker not in speakers:
                    speakers[speaker] = {
                        'words': [],
                        'total_confidence': 0,
                        'word_count': 0,
                        'speaking_time': 0
                    }
                speakers[speaker]['words'].append(word)
                speakers[speaker]['total_confidence'] += word.get('confidence', 0)
                speakers[speaker]['word_count'] += 1
        
        # Calculate average confidence and speaking patterns for each speaker
        speaker_profiles = {}
        for speaker_id, data in speakers.items():
            avg_confidence = data['total_confidence'] / data['word_count'] if data['word_count'] > 0 else 0
            
            # Determine speaking characteristics
            first_words = ' '.join([w['text'] for w in data['words'][:20]])  # First 20 words
            is_chair_like = any(phrase in first_words.lower() for phrase in 
                              ['call to order', 'meeting to order', 'thank you all', 'welcome'])
            
            speaker_profiles[speaker_id] = {
                'avg_confidence': avg_confidence,
                'word_count': data['word_count'],
                'is_chair_like': is_chair_like,
                'assigned_name': None,
                'assignment_confidence': 0
            }
        
        # Attempt to match speakers to committee members
        if committee_members:
            # Sort speakers by word count (chair usually speaks most)
            sorted_speakers = sorted(speaker_profiles.items(), 
                                   key=lambda x: (x[1]['is_chair_like'], x[1]['word_count']), 
                                   reverse=True)
            
            # Simple assignment logic (can be enhanced with voice analysis)
            available_members = committee_members.copy()
            
            for i, (speaker_id, profile) in enumerate(sorted_speakers[:len(committee_members)]):
                if available_members:
                    # Assign most likely member based on speaking pattern
                    assigned_member = available_members.pop(0)
                    
                    # Calculate confidence based on speaking patterns and audio quality
                    base_confidence = profile['avg_confidence']
                    pattern_bonus = 20 if profile['is_chair_like'] and i == 0 else 0  # Chair bonus
                    word_count_bonus = min(10, profile['word_count'] / 50)  # More words = higher confidence
                    
                    assignment_confidence = min(100, base_confidence + pattern_bonus + word_count_bonus)
                    
                    speaker_profiles[speaker_id]['assigned_name'] = assigned_member
                    speaker_profiles[speaker_id]['assignment_confidence'] = assignment_confidence
                    
                    logger.info(f"Assigned {speaker_id} to {assigned_member} with {assignment_confidence:.1f}% confidence")
        
        # Apply speaker assignments to transcript
        for word in transcript_data['words']:
            speaker_id = word.get('speaker')
            if speaker_id and speaker_id in speaker_profiles:
                profile = speaker_profiles[speaker_id]
                
                if profile['assigned_name'] and profile['assignment_confidence'] >= confidence_threshold:
                    # High confidence - use real name
                    word['speaker'] = profile['assigned_name']
                elif profile['assigned_name'] and profile['assignment_confidence'] >= 40:
                    # Medium confidence - add uncertainty indicator
                    word['speaker'] = f"{profile['assigned_name']} (uncertain)"
                else:
                    # Low confidence - use descriptive label
                    if profile['is_chair_like']:
                        word['speaker'] = "Chair"
                    elif profile['word_count'] > 100:
                        word['speaker'] = "Presenter"
                    else:
                        word['speaker'] = f"Unknown Speaker"
        
        return transcript_data
        
    except Exception as e:
        logger.error(f"Error identifying speakers: {e}")
        return transcript_data
    """Create a folder in Nextcloud using WebDAV."""
    try:
        webdav_url = f"{NEXTCLOUD_URL}/remote.php/dav/files/{NEXTCLOUD_USERNAME}/{folder_path}"
        
        response = requests.request(
            'MKCOL',
            webdav_url,
            auth=(NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN),
            timeout=30
        )
        
        # 201 = created, 405 = already exists
        if response.status_code in [201, 405]:
            logger.info(f"Nextcloud folder ready: {folder_path}")
            return True
        else:
            logger.error(f"Failed to create Nextcloud folder: {response.status_code} - {response.text}")
            return False
            
    except Exception as e:
        logger.error(f"Error creating Nextcloud folder: {e}")
        return False


def upload_to_nextcloud(file_path, nextcloud_path):
    """Upload a file to Nextcloud using WebDAV."""
    try:
        webdav_url = f"{NEXTCLOUD_URL}/remote.php/dav/files/{NEXTCLOUD_USERNAME}/{nextcloud_path}"
        
        with open(file_path, 'rb') as f:
            response = requests.put(
                webdav_url,
                data=f,
                auth=(NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN),
                timeout=300  # 5 minutes for file upload
            )
        
        if response.status_code in [201, 204]:
            logger.info(f"Successfully uploaded to Nextcloud: {nextcloud_path}")
            return True
        else:
            logger.error(f"Failed to upload to Nextcloud: {response.status_code} - {response.text}")
            return False
            
    except Exception as e:
        logger.error(f"Error uploading to Nextcloud: {e}")
        return False


def get_nextcloud_share_link(nextcloud_path):
    """Create a public share link for a Nextcloud file with edit permissions."""
    try:
        # Create share via OCS API
        share_url = f"{NEXTCLOUD_URL}/ocs/v2.php/apps/files_sharing/api/v1/shares"
        
        data = {
            'path': f"/{nextcloud_path}",
            'shareType': 3,  # Public link
            'permissions': 15  # Full permissions (read + write + create + delete + share)
        }
        
        headers = {
            'OCS-APIRequest': 'true',
            'Content-Type': 'application/x-www-form-urlencoded'
        }
        
        response = requests.post(
            share_url,
            data=data,
            headers=headers,
            auth=(NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN),
            timeout=30
        )
        
        if response.status_code == 200:
            # Parse XML response to get share URL
            from xml.etree import ElementTree as ET
            root = ET.fromstring(response.text)
            
            url_element = root.find('.//url')
            if url_element is not None:
                share_link = url_element.text
                logger.info(f"Created Nextcloud editable share link: {share_link}")
                return share_link
            else:
                logger.error("Could not find URL in share response")
                return None
        else:
            logger.error(f"Failed to create share: {response.status_code} - {response.text}")
            return None
            
    except Exception as e:
        logger.error(f"Error creating Nextcloud share link: {e}")
        return None


def get_clean_committee_mapping():
    """Return mapping of committee names for clean folder/file naming."""
    return {
        'interim': {
            'capitol buildings planning': 'Capitol Buildings Planning',
            'capitol security': 'Capitol Security',
            'courts corrections justice': 'Courts, Corrections & Justice',
            'courts, corrections & justice': 'Courts, Corrections & Justice',
            'economic rural development': 'Economic & Rural Development & Policy',
            'economic & rural development': 'Economic & Rural Development & Policy',
            'facilities review': 'Facilities Review',
            'federal funding stabilization': 'Federal Funding Stabilization',
            'indian affairs': 'Indian Affairs',
            'interim legislative ethics': 'Interim Legislative Ethics',
            'legislative ethics': 'Interim Legislative Ethics',
            'investments pensions oversight': 'Investments & Pensions Oversight',
            'investments & pensions oversight': 'Investments & Pensions Oversight',
            'land grant': 'Land Grant',
            'legislative council': 'Legislative Council',
            'legislative education study': 'Legislative Education Study',
            'legislative finance': 'Legislative Finance',
            'legislative interim committee working group': 'Legislative Interim Committee Working Group',
            'legislative health human services': 'Legislative Health & Human Services',
            'legislative health & human services': 'Legislative Health & Human Services',
            'military veterans affairs': 'Military & Veterans\' Affairs',
            'military & veterans\' affairs': 'Military & Veterans\' Affairs',
            'mortgage finance authority act oversight': 'Mortgage Finance Authority Act Oversight',
            'mortgage finance authority oversight': 'Mortgage Finance Authority Act Oversight',
            'new mexico finance authority oversight': 'New Mexico Finance Authority Oversight',
            'public school capital outlay oversight': 'Public School Capital Outlay Oversight',
            'public school capital outlay council': 'Public School Capital Outlay Oversight',
            'public school capital outlay oversight task': 'Public School Capital Outlay Oversight',
            'radioactive hazardous materials': 'Radioactive & Hazardous Materials',
            'radioactive & hazardous materials': 'Radioactive & Hazardous Materials',
            'revenue stabilization tax policy': 'Revenue Stabilization & Tax Policy',
            'revenue stabilization & tax policy': 'Revenue Stabilization & Tax Policy',
            'science technology telecommunications': 'Science, Technology & Telecommunications',
            'science, technology & telecommunications': 'Science, Technology & Telecommunications',
            'tobacco settlement revenue oversight': 'Tobacco Settlement Revenue Oversight',
            'transportation infrastructure revenue': 'Transportation Infrastructure Revenue',
            'transportation infrastructure revenue subcommittee': 'Transportation Infrastructure Revenue',
            'water natural resources': 'Water & Natural Resources',
            'water & natural resources': 'Water & Natural Resources'
        },
        'house': {
            'agriculture acequias water resources': 'Agriculture, Acequias And Water Resources',
            'agriculture, acequias and water resources': 'Agriculture, Acequias And Water Resources',
            'appropriations finance': 'Appropriations & Finance',
            'appropriations & finance': 'Appropriations & Finance',
            'commerce economic development': 'Commerce & Economic Development Committee',
            'commerce & economic development committee': 'Commerce & Economic Development Committee',
            'consumer public affairs': 'Consumer & Public Affairs',
            'consumer & public affairs': 'Consumer & Public Affairs',
            'education': 'Education',
            'energy environment natural resources': 'Energy, Environment & Natural Resources',
            'energy, environment & natural resources': 'Energy, Environment & Natural Resources',
            'government elections indian affairs': 'Government, Elections & Indian Affairs',
            'government, elections & indian affairs': 'Government, Elections & Indian Affairs',
            'health human services': 'Health & Human Services',
            'health & human services': 'Health & Human Services',
            'judiciary': 'Judiciary',
            'labor veterans military affairs': 'Labor, Veterans\' And Military Affairs Committee',
            'labor, veterans\' and military affairs committee': 'Labor, Veterans\' And Military Affairs Committee',
            'rules order business': 'Rules & Order Of Business',
            'rules & order of business': 'Rules & Order Of Business',
            'rural development land grants cultural affairs': 'Rural Development, Land Grants And Cultural Affairs',
            'rural development, land grants and cultural affairs': 'Rural Development, Land Grants And Cultural Affairs',
            'taxation revenue': 'Taxation & Revenue',
            'taxation & revenue': 'Taxation & Revenue',
            'transportation public works capital improvements': 'Transportation, Public Works & Capital Improvements',
            'transportation, public works & capital improvements': 'Transportation, Public Works & Capital Improvements'
        },
        'senate': {
            'committees committee': 'Committees\' Committee',
            'committees\' committee': 'Committees\' Committee',
            'conservation': 'Conservation',
            'education': 'Education',
            'finance': 'Finance',
            'health public affairs': 'Health & Public Affairs',
            'health & public affairs': 'Health & Public Affairs',
            'indian rural cultural affairs': 'Indian, Rural & Cultural Affairs',
            'indian, rural & cultural affairs': 'Indian, Rural & Cultural Affairs',
            'judiciary': 'Judiciary',
            'rules': 'Rules',
            'tax business transportation': 'Tax, Business & Transportation',
            'tax, business & transportation': 'Tax, Business & Transportation'
        }
    }


def parse_meeting_title(title):
    """Parse meeting title to determine folder structure and clean title."""
    try:
        # Clean up the title first - remove extra whitespace and newlines
        clean_title = re.sub(r'\s+', ' ', title).strip()
        
        # Split on the first " - " to get prefix and committee name
        if ' - ' in clean_title:
            prefix, committee_text = clean_title.split(' - ', 1)
            prefix = prefix.strip().upper()
            committee_text = committee_text.strip()
            
            # Get committee mapping
            committee_mapping = get_clean_committee_mapping()
            
            # Clean committee text for matching (remove special chars, normalize)
            committee_for_matching = re.sub(r'[^\w\s]', ' ', committee_text.lower())
            committee_for_matching = re.sub(r'\s+', ' ', committee_for_matching).strip()
            
            # Try to find a matching committee
            clean_committee_name = None
            committee_type = None
            
            if prefix == 'IC':
                # Interim Committee
                committee_type = 'interim'
                for key, value in committee_mapping['interim'].items():
                    if key in committee_for_matching or committee_for_matching in key:
                        clean_committee_name = value
                        break
                
                if not clean_committee_name:
                    # Fallback - try partial matches
                    for key, value in committee_mapping['interim'].items():
                        key_words = set(key.split())
                        text_words = set(committee_for_matching.split())
                        if len(key_words.intersection(text_words)) >= 2:  # At least 2 words match
                            clean_committee_name = value
                            break
                
                if clean_committee_name:
                    return {
                        'type': 'interim',
                        'folder_path': f"Legislative Transcription/Interim/{clean_committee_name}",
                        'committee_name': clean_committee_name,
                        'chamber': None,
                        'clean_title': f"IC - {clean_committee_name}"
                    }
                    
            elif prefix in ['HOUSE', 'SENATE']:
                # Legislative Session Committee
                chamber = prefix.title()
                committee_type = chamber.lower()
                
                for key, value in committee_mapping[committee_type].items():
                    if key in committee_for_matching or committee_for_matching in key:
                        clean_committee_name = value
                        break
                
                if not clean_committee_name:
                    # Fallback - try partial matches
                    for key, value in committee_mapping[committee_type].items():
                        key_words = set(key.split())
                        text_words = set(committee_for_matching.split())
                        if len(key_words.intersection(text_words)) >= 2:  # At least 2 words match
                            clean_committee_name = value
                            break
                
                if clean_committee_name:
                    return {
                        'type': 'session',
                        'folder_path': f"Legislative Transcription/LegSession/{chamber}/{clean_committee_name}",
                        'committee_name': clean_committee_name,
                        'chamber': chamber,
                        'clean_title': f"{chamber} - {clean_committee_name}"
                    }
            
            # If no clean match found, log and use fallback
            logger.warning(f"Could not match committee '{committee_text}' to known committees")
            fallback_name = committee_text[:50]  # Limit length
            return {
                'type': 'unknown',
                'folder_path': f"Legislative Transcription/Other/{fallback_name}",
                'committee_name': fallback_name,
                'chamber': None,
                'clean_title': f"{prefix} - {fallback_name}"
            }
        else:
            # No " - " separator found, use default structure
            logger.warning(f"Could not parse meeting title format: {clean_title[:50]}...")
            clean_committee = re.sub(r'[^\w\s-]', '', clean_title)
            clean_committee = re.sub(r'\s+', ' ', clean_committee).strip()[:50]
            
            return {
                'type': 'unknown',
                'folder_path': f"Legislative Transcription/Other/{clean_committee or 'Unknown'}",
                'committee_name': clean_committee or 'Unknown Committee',
                'chamber': None,
                'clean_title': clean_title
            }
            
    except Exception as e:
        logger.error(f"Error parsing meeting title: {e}")
        return {
            'type': 'error',
            'folder_path': 'Legislative Transcription/Other/Parse_Error',
            'committee_name': 'Unknown Committee',
            'chamber': None,
            'clean_title': title
        }


def create_folder_hierarchy(base_folder_path, date_folder):
    """Create the full folder hierarchy including date subfolder."""
    try:
        # Create base folder structure first
        folder_parts = base_folder_path.split('/')
        current_path = ""
        
        for part in folder_parts:
            if current_path:
                current_path += f"/{part}"
            else:
                current_path = part
            
            success = create_nextcloud_folder(current_path)
            if not success:
                logger.warning(f"Could not create folder: {current_path}")
        
        # Create date subfolder
        full_path = f"{base_folder_path}/{date_folder}"
        success = create_nextcloud_folder(full_path)
        
        return success, full_path
        
    except Exception as e:
        logger.error(f"Error creating folder hierarchy: {e}")
        return False, None
    """Parse meeting title to determine folder structure and clean title."""
    try:
        # Clean up the title first - remove extra whitespace and newlines
        clean_title = re.sub(r'\s+', ' ', title).strip()
        
        # Split on the first " - " to get prefix and committee name
        if ' - ' in clean_title:
            prefix, committee_name = clean_title.split(' - ', 1)
            prefix = prefix.strip().upper()
            committee_name = committee_name.strip()
            
            # Clean committee name for folder use (remove special chars, keep spaces)
            clean_committee = re.sub(r'[^\w\s-]', '', committee_name)
            clean_committee = re.sub(r'\s+', ' ', clean_committee).strip()
            
            if prefix == 'IC':
                # Interim Committee
                folder_type = 'Interim'
                chamber = None
                return {
                    'type': 'interim',
                    'folder_path': f"Legislative Transcription/Interim/{clean_committee}",
                    'committee_name': committee_name,
                    'chamber': None,
                    'clean_title': clean_title
                }
            elif prefix in ['HOUSE', 'SENATE']:
                # Legislative Session Committee
                folder_type = 'LegSession'
                chamber = prefix.title()  # House or Senate
                return {
                    'type': 'session',
                    'folder_path': f"Legislative Transcription/LegSession/{chamber}/{clean_committee}",
                    'committee_name': committee_name,
                    'chamber': chamber,
                    'clean_title': clean_title
                }
            else:
                # Unknown prefix, fall back to original structure
                logger.warning(f"Unknown meeting prefix '{prefix}', using default folder structure")
                return {
                    'type': 'unknown',
                    'folder_path': f"Legislative Transcription/Other/{clean_committee or 'Unknown'}",
                    'committee_name': committee_name or 'Unknown Committee',
                    'chamber': None,
                    'clean_title': clean_title
                }
        else:
            # No " - " separator found, use default structure
            logger.warning(f"Could not parse meeting title format, using default folder structure: {clean_title[:50]}...")
            clean_committee = re.sub(r'[^\w\s-]', '', clean_title)
            clean_committee = re.sub(r'\s+', ' ', clean_committee).strip()[:50]  # Limit length
            
            return {
                'type': 'unknown',
                'folder_path': f"Legislative Transcription/Other/{clean_committee or 'Unknown'}",
                'committee_name': clean_committee or 'Unknown Committee',
                'chamber': None,
                'clean_title': clean_title
            }
            
    except Exception as e:
        logger.error(f"Error parsing meeting title: {e}")
        return {
            'type': 'error',
            'folder_path': 'Legislative Transcription/Other/Parse_Error',
            'committee_name': 'Unknown Committee',
            'chamber': None,
            'clean_title': title
        }


def create_folder_hierarchy(base_folder_path, date_folder):
    """Create the full folder hierarchy including date subfolder."""
    try:
        # Create base folder structure first
        folder_parts = base_folder_path.split('/')
        current_path = ""
        
        for part in folder_parts:
            if current_path:
                current_path += f"/{part}"
            else:
                current_path = part
            
            success = create_nextcloud_folder(current_path)
            if not success:
                logger.warning(f"Could not create folder: {current_path}")
        
        # Create date subfolder
        full_path = f"{base_folder_path}/{date_folder}"
        success = create_nextcloud_folder(full_path)
        
        return success, full_path
        
    except Exception as e:
        logger.error(f"Error creating folder hierarchy: {e}")
        return False, None


def check_nextcloud_file_exists(nextcloud_path):
    """Check if a file exists in Nextcloud using WebDAV."""
    try:
        webdav_url = f"{NEXTCLOUD_URL}/remote.php/dav/files/{NEXTCLOUD_USERNAME}/{nextcloud_path}"
        
        response = requests.head(
            webdav_url,
            auth=(NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN),
            timeout=30
        )
        
        # 200 = file exists, 404 = file doesn't exist
        return response.status_code == 200
        
    except Exception as e:
        logger.error(f"Error checking if Nextcloud file exists: {e}")
        return False


def generate_unique_filename(base_filename, folder_path):
    """Generate a unique filename if conflicts exist in Nextcloud."""
    try:
        # Split filename and extension
        name_part, ext = os.path.splitext(base_filename)
        counter = 1
        
        # Check if base filename exists
        test_path = f"{folder_path}/{base_filename}"
        if not check_nextcloud_file_exists(test_path):
            return base_filename
        
        # Try numbered versions until we find one that doesn't exist
        while counter <= 50:  # Reasonable limit to prevent infinite loops
            new_filename = f"{name_part}_v{counter}{ext}"
            test_path = f"{folder_path}/{new_filename}"
            
            if not check_nextcloud_file_exists(test_path):
                logger.info(f"Generated unique filename: {new_filename}")
                return new_filename
            
            counter += 1
        
        # If we still have conflicts after 50 tries, add timestamp
        timestamp = datetime.datetime.now().strftime('%H%M%S')
        unique_filename = f"{name_part}_{timestamp}{ext}"
        logger.warning(f"Used timestamp for unique filename: {unique_filename}")
        return unique_filename
        
    except Exception as e:
        logger.error(f"Error generating unique filename: {e}")
        # Fallback to timestamp-based name
        timestamp = datetime.datetime.now().strftime('%H%M%S')
        name_part, ext = os.path.splitext(base_filename)
        return f"{name_part}_{timestamp}{ext}"


def extract_meeting_datetime(title):
    """Extract meeting date and time from the title text."""
    try:
        # Common date patterns to look for
        date_patterns = [
            # "Jun 25 2025", "June 25, 2025", etc.
            r'([A-Za-z]{3,9})\s+(\d{1,2}),?\s+(\d{4})',
            # "2025-06-25", "06/25/2025", etc.
            r'(\d{4})[-/](\d{1,2})[-/](\d{1,2})',
            r'(\d{1,2})[-/](\d{1,2})[-/](\d{4})'
        ]
        
        # Time patterns to look for
        time_patterns = [
            # "916 AM-1158 AM", "9:16 AM-11:58 AM", etc.
            r'(\d{1,2}):?(\d{2})\s*(AM|PM)[-–]\s*(\d{1,2}):?(\d{2})\s*(AM|PM)',
            # Single time "916 AM", "9:16 AM"
            r'(\d{1,2}):?(\d{2})\s*(AM|PM)'
        ]
        
        meeting_date = None
        meeting_time = None
        
        # Extract date
        for pattern in date_patterns:
            match = re.search(pattern, title, re.IGNORECASE)
            if match:
                groups = match.groups()
                if len(groups) == 3:
                    if groups[0].isdigit():  # Format: YYYY-MM-DD or MM/DD/YYYY
                        if len(groups[0]) == 4:  # YYYY-MM-DD
                            year, month, day = groups
                        else:  # MM/DD/YYYY
                            month, day, year = groups
                    else:  # Format: Month DD YYYY
                        month_name, day, year = groups
                        # Convert month name to number
                        month_mapping = {
                            'jan': '01', 'january': '01',
                            'feb': '02', 'february': '02',
                            'mar': '03', 'march': '03',
                            'apr': '04', 'april': '04',
                            'may': '05', 'may': '05',
                            'jun': '06', 'june': '06',
                            'jul': '07', 'july': '07',
                            'aug': '08', 'august': '08',
                            'sep': '09', 'september': '09',
                            'oct': '10', 'october': '10',
                            'nov': '11', 'november': '11',
                            'dec': '12', 'december': '12'
                        }
                        month = month_mapping.get(month_name.lower()[:3])
                        if not month:
                            continue
                    
                    # Format as MMDDYY
                    try:
                        month_num = int(month)
                        day_num = int(day)
                        year_num = int(year)
                        
                        # Convert to MMDDYY format
                        meeting_date = f"{month_num:02d}{day_num:02d}{year_num % 100:02d}"
                        break
                    except ValueError:
                        continue
        
        # Extract time
        for pattern in time_patterns:
            match = re.search(pattern, title, re.IGNORECASE)
            if match:
                groups = match.groups()
                if len(groups) >= 3:  # At least start time
                    if len(groups) == 6:  # Start and end time
                        start_hour, start_min, start_period, end_hour, end_min, end_period = groups
                        meeting_time = f"{start_hour}{start_min} {start_period.upper()}-{end_hour}{end_min} {end_period.upper()}"
                    else:  # Single time
                        hour, minute, period = groups[:3]
                        meeting_time = f"{hour}{minute} {period.upper()}"
                    break
        
        return meeting_date, meeting_time
        
    except Exception as e:
        logger.warning(f"Error extracting meeting date/time: {e}")
        return None, None


def save_transcript_to_nextcloud(transcript_data, title, entry_url):
    """Save transcript as DOCX to Nextcloud with structured folder organization."""
    try:
        if not all([NEXTCLOUD_URL, NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN]):
            logger.error("Nextcloud configuration missing")
            return None
        
        # Parse the meeting title to determine folder structure
        meeting_info = parse_meeting_title(title)
        
        # Extract meeting date and time from title
        meeting_date, meeting_time = extract_meeting_datetime(title)
        
        # Use extracted date for folder, fallback to current date
        if meeting_date:
            # Convert MMDDYY to YYYY-MM-DD for folder structure
            mm = meeting_date[:2]
            dd = meeting_date[2:4]
            yy = meeting_date[4:6]
            # Assume 20xx for years
            yyyy = f"20{yy}" if int(yy) < 50 else f"19{yy}"
            date_folder = f"{yyyy}-{mm}-{dd}"
            logger.info(f"Using extracted meeting date: {date_folder}")
        else:
            date_folder = datetime.datetime.now().strftime('%Y-%m-%d')
            logger.warning(f"Could not extract meeting date from title, using current date: {date_folder}")
        
        # Create the full folder hierarchy
        success, full_folder_path = create_folder_hierarchy(meeting_info['folder_path'], date_folder)
        
        if not success:
            logger.error("Failed to create folder hierarchy, falling back to simple structure")
            # Fallback to simple structure
            simple_folder = f"Legislative Transcription/{date_folder}"
            create_nextcloud_folder("Legislative Transcription")
            create_nextcloud_folder(simple_folder)
            full_folder_path = simple_folder
        
        # Create base filename using extracted date and time
        if meeting_date:
            date_prefix = meeting_date  # Already in MMDDYY format
        else:
            # Fallback to current date in MMDDYY format
            now = datetime.datetime.now()
            date_prefix = f"{now.month:02d}{now.day:02d}{now.year % 100:02d}"
        
        # Use clean committee abbreviation from parsing
        committee_name = meeting_info['committee_name']  # This will be the abbreviation like "LHHS"
        
        # Build filename components
        filename_parts = [date_prefix, committee_name]
        
        if meeting_time:
            # Clean up time format - remove spaces and standardize
            clean_time = meeting_time.replace(' ', '').replace('-', '-')  # Remove all spaces
            filename_parts.append(clean_time)
            logger.info(f"Using extracted meeting time: {clean_time}")
        else:
            logger.warning("Could not extract meeting time from title")
        
        # Join parts with dashes and clean for filename
        filename_base = "-".join(filename_parts)
        # Clean for filesystem (remove problematic characters)
        filename_base = re.sub(r'[<>:"/\\|?*]', '', filename_base)
        
        base_filename = f"{filename_base}.docx"
        
        # Generate unique filename if conflicts exist
        unique_filename = generate_unique_filename(base_filename, full_folder_path)
        
        # Create the DOCX document with enhanced metadata
        doc = create_docx_document_with_metadata(transcript_data, title, entry_url, meeting_info)
        if not doc:
            logger.error("Failed to create DOCX document")
            return None
        
        # Save to temporary file
        temp_path = os.path.join(DOWNLOAD_DIR, unique_filename)
        doc.save(temp_path)
        
        # Upload to Nextcloud
        nextcloud_file_path = f"{full_folder_path}/{unique_filename}"
        
        if upload_to_nextcloud(temp_path, nextcloud_file_path):
            logger.info(f"Successfully saved transcript to: {nextcloud_file_path}")
            
            # Create share link
            share_link = get_nextcloud_share_link(nextcloud_file_path)
            
            # Clean up temporary file
            try:
                os.remove(temp_path)
            except Exception as e:
                logger.warning(f"Could not remove temporary file: {e}")
            
            return {
                'nextcloud_path': nextcloud_file_path,
                'share_link': share_link,
                'filename': unique_filename,
                'meeting_info': meeting_info,
                'folder_path': full_folder_path,
                'meeting_date': meeting_date,
                'meeting_time': meeting_time
            }
        else:
            logger.error("Failed to upload to Nextcloud")
            return None
            
    except Exception as e:
        logger.error(f"Error saving transcript to Nextcloud: {e}")
        return None


def create_docx_document_with_metadata(transcript_data, title, entry_url, meeting_info):
    """Create a DOCX document with enhanced meeting metadata."""
    try:
        # Create a new document
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add meeting information section
        doc.add_heading('Meeting Information', level=1)
        
        # Committee information
        if meeting_info['type'] == 'interim':
            info_paragraph = doc.add_paragraph()
            info_paragraph.add_run('Meeting Type: ').bold = True
            info_paragraph.add_run('Interim Committee')
            
        elif meeting_info['type'] == 'session':
            info_paragraph = doc.add_paragraph()
            info_paragraph.add_run('Meeting Type: ').bold = True
            info_paragraph.add_run('Legislative Session Committee')
            
            info_paragraph = doc.add_paragraph()
            info_paragraph.add_run('Chamber: ').bold = True
            info_paragraph.add_run(meeting_info['chamber'])
        
        info_paragraph = doc.add_paragraph()
        info_paragraph.add_run('Committee: ').bold = True
        info_paragraph.add_run(meeting_info['committee_name'])
        
        # Standard metadata
        info_paragraph = doc.add_paragraph()
        info_paragraph.add_run('Source URL: ').bold = True
        info_paragraph.add_run(entry_url)
        
        info_paragraph = doc.add_paragraph()
        info_paragraph.add_run('Transcribed by: ').bold = True
        info_paragraph.add_run('LWE.Vote')
        
        info_paragraph = doc.add_paragraph()
        info_paragraph.add_run('Transcribed on: ').bold = True
        info_paragraph.add_run(datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        
        if transcript_data:
            info_paragraph = doc.add_paragraph()
            info_paragraph.add_run('Audio duration: ').bold = True
            info_paragraph.add_run(f"{transcript_data.get('audio_duration', 'Unknown')} seconds")
            
            info_paragraph = doc.add_paragraph()
            info_paragraph.add_run('Confidence score: ').bold = True
            info_paragraph.add_run(f"{transcript_data.get('confidence', 'Unknown')}")
        
        # Add separator
        doc.add_paragraph('_' * 80)
        
        # Add transcript section
        doc.add_heading('Transcript', level=1)
        
        # Format and add transcript content
        if transcript_data and "words" in transcript_data:
            current_speaker = None
            current_paragraph = None
            
            for word in transcript_data["words"]:
                speaker = word.get("speaker")
                
                if speaker != current_speaker:
                    # New speaker, create new paragraph
                    current_paragraph = doc.add_paragraph()
                    current_speaker = speaker
                    
                    if speaker is not None:
                        speaker_run = current_paragraph.add_run(f"Speaker {speaker}: ")
                        speaker_run.bold = True
                        current_paragraph.add_run(f"{word['text']} ")
                    else:
                        current_paragraph.add_run(f"{word['text']} ")
                else:
                    # Same speaker, continue paragraph
                    if current_paragraph:
                        current_paragraph.add_run(f"{word['text']} ")
        
        elif transcript_data and transcript_data.get("text"):
            # Fallback to plain text if words are not available
            doc.add_paragraph(transcript_data["text"])
        else:
            doc.add_paragraph("No transcription text available.")
        
        return doc
        
    except Exception as e:
        logger.error(f"Error creating DOCX document with metadata: {e}")
        return None


def get_current_entries():
    """Fetch and parse the current entries from the website using proxy."""
    try:
        # Use proxy-enabled request
        response = make_request_with_proxy(URL)
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        entries = []
        
        for entry_element in soup.select('table'):
            entry_text = entry_element.get_text().strip()
            
            parent_a = entry_element.find_parent('a')
            entry_link = None
            if parent_a and 'href' in parent_a.attrs:
                entry_link = parent_a['href']
                if entry_link and not entry_link.startswith('http'):
                    base_url = '/'.join(URL.split('/')[:3])
                    entry_link = f"{base_url}{entry_link if entry_link.startswith('/') else '/' + entry_link}"
            
            if entry_text:
                entries.append({
                    'text': entry_text,
                    'link': entry_link,
                    'timestamp': datetime.datetime.now().isoformat()
                })
        
        return entries
    
    except Exception as e:
        logger.error(f"Error fetching entries: {e}")
        return []


def read_stored_entries():
    """Read the previously stored entries."""
    if not os.path.exists(ENTRIES_FILE):
        return []
    
    try:
        with open(ENTRIES_FILE, 'r') as f:
            return json.load(f)
    except json.JSONDecodeError:
        logger.error(f"Error parsing {ENTRIES_FILE}, starting with empty entries")
        return []


def write_entries(entries):
    """Write entries to the storage file."""
    with open(ENTRIES_FILE, 'w') as f:
        json.dump(entries, f, indent=2)


def read_processed_entries():
    """Read the previously processed entries with error handling for format changes."""
    if not os.path.exists(PROCESSED_ENTRIES_FILE):
        return []
    
    try:
        with open(PROCESSED_ENTRIES_FILE, 'r') as f:
            content = f.read().strip()
            if not content:
                return []
            
            entries = json.loads(content)
            
            # Validate that entries is a list
            if not isinstance(entries, list):
                logger.error(f"Processed entries file contains invalid format (not a list), resetting...")
                return []
            
            # Check if entries have expected structure and clean up if needed
            valid_entries = []
            for entry in entries:
                if isinstance(entry, dict) and 'text' in entry:
                    # Convert old Google Docs format to new Nextcloud format if needed
                    if 'transcription' in entry:
                        transcription = entry['transcription']
                        if isinstance(transcription, dict):
                            # Update old google_doc_url to nextcloud_result format
                            if 'google_doc_url' in transcription and 'nextcloud_result' not in transcription:
                                transcription['nextcloud_result'] = {
                                    'legacy_google_doc': transcription.pop('google_doc_url'),
                                    'migrated_to_nextcloud': True
                                }
                    
                    valid_entries.append(entry)
                else:
                    logger.warning(f"Skipping invalid entry format: {entry}")
            
            return valid_entries
            
    except json.JSONDecodeError as e:
        logger.error(f"Error parsing {PROCESSED_ENTRIES_FILE}: {e}")
        logger.info(f"Backing up corrupted file and starting fresh...")
        
        # Backup the corrupted file
        backup_file = f"{PROCESSED_ENTRIES_FILE}.backup.{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
        try:
            os.rename(PROCESSED_ENTRIES_FILE, backup_file)
            logger.info(f"Backed up corrupted file to: {backup_file}")
        except Exception as backup_error:
            logger.error(f"Could not backup corrupted file: {backup_error}")
        
        return []
    except Exception as e:
        logger.error(f"Unexpected error reading processed entries: {e}")
        return []


def write_processed_entry(entry):
    """Add an entry to the processed entries file."""
    processed = read_processed_entries()
    processed.append(entry)
    
    with open(PROCESSED_ENTRIES_FILE, 'w') as f:
        json.dump(processed, f, indent=2)


def format_meeting_date_for_email(meeting_date):
    """Format meeting date for email display (e.g., Tuesday, July 1, 2025)."""
    try:
        if not meeting_date or len(meeting_date) != 6:
            return datetime.datetime.now().strftime('%A, %B %d, %Y')
        
        # Parse MMDDYY format
        mm = int(meeting_date[:2])
        dd = int(meeting_date[2:4])
        yy = int(meeting_date[4:6])
        
        # Assume 20xx for years
        yyyy = 2000 + yy if yy < 50 else 1900 + yy
        
        # Create date object and format
        date_obj = datetime.datetime(yyyy, mm, dd)
        return date_obj.strftime('%A, %B %d, %Y')
        
    except Exception as e:
        logger.warning(f"Error formatting meeting date: {e}")
        return datetime.datetime.now().strftime('%A, %B %d, %Y')


def format_meeting_date_for_subject(meeting_date):
    """Format meeting date for email subject (e.g., 7/1/25)."""
    try:
        if not meeting_date or len(meeting_date) != 6:
            now = datetime.datetime.now()
            return f"{now.month}/{now.day}/{now.year % 100}"
        
        # Parse MMDDYY format
        mm = int(meeting_date[:2])
        dd = int(meeting_date[2:4])
        yy = int(meeting_date[4:6])
        
        return f"{mm}/{dd}/{yy}"
        
    except Exception as e:
        logger.warning(f"Error formatting meeting date for subject: {e}")
        now = datetime.datetime.now()
        return f"{now.month}/{now.day}/{now.year % 100}"


def send_notification(new_entries, transcript_results=None):
    """Send HTML email notification about new entries with Nextcloud links."""
    if not all([EMAIL_USER, EMAIL_PASSWORD]):
        logger.warning("Email configuration missing. Skipping notification.")
        return
    
    try:
        for recipient in EMAIL_RECIPIENTS:
            for i, entry in enumerate(new_entries):
                # Get transcript result for this entry
                result = transcript_results[i] if transcript_results and i < len(transcript_results) else None
                
                if not result:
                    logger.warning(f"No transcript result for entry {i}, skipping email")
                    continue
                
                # Extract meeting information
                meeting_info = result.get('meeting_info', {})
                committee_name = meeting_info.get('committee_name', 'Unknown Committee')
                
                # Get meeting date and time from result
                meeting_date = result.get('meeting_date')
                meeting_time = result.get('meeting_time', 'Time not available')
                
                # Format dates for email
                formatted_date = format_meeting_date_for_email(meeting_date)
                subject_date = format_meeting_date_for_subject(meeting_date)
                
                # Create email
                msg = MIMEMultipart('alternative')
                msg['From'] = EMAIL_USER
                msg['To'] = recipient
                msg['Subject'] = f"New Legislature Transcript - {committee_name} - {subject_date}"
                
                # HTML email template
                html_template = '''<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.0 Transitional//EN" "http://www.w3.org/TR/xhtml1/DTD/xhtml1-transitional.dtd">
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office" lang="en">
<head>
<title></title>
<meta charset="UTF-8" />
<meta http-equiv="Content-Type" content="text/html; charset=UTF-8" />
<!--[if !mso]>-->
<meta http-equiv="X-UA-Compatible" content="IE=edge" />
<!--<![endif]-->
<meta name="x-apple-disable-message-reformatting" content="" />
<meta content="target-densitydpi=device-dpi" name="viewport" />
<meta content="true" name="HandheldFriendly" />
<meta content="width=device-width" name="viewport" />
<meta name="format-detection" content="telephone=no, date=no, address=no, email=no, url=no" />
<style type="text/css">
table {
border-collapse: separate;
table-layout: fixed;
mso-table-lspace: 0pt;
mso-table-rspace: 0pt
}
table td {
border-collapse: collapse
}
.ExternalClass {
width: 100%
}
.ExternalClass,
.ExternalClass p,
.ExternalClass span,
.ExternalClass font,
.ExternalClass td,
.ExternalClass div {
line-height: 100%
}
body, a, li, p, h1, h2, h3 {
-ms-text-size-adjust: 100%;
-webkit-text-size-adjust: 100%;
}
html {
-webkit-text-size-adjust: none !important
}
body, #innerTable {
-webkit-font-smoothing: antialiased;
-moz-osx-font-smoothing: grayscale
}
#innerTable img+div {
display: none;
display: none !important
}
img {
Margin: 0;
padding: 0;
-ms-interpolation-mode: bicubic
}
h1, h2, h3, p, a {
line-height: inherit;
overflow-wrap: normal;
white-space: normal;
word-break: break-word
}
a {
text-decoration: none
}
h1, h2, h3, p {
min-width: 100%!important;
width: 100%!important;
max-width: 100%!important;
display: inline-block!important;
border: 0;
padding: 0;
margin: 0
}
a[x-apple-data-detectors] {
color: inherit !important;
text-decoration: none !important;
font-size: inherit !important;
font-family: inherit !important;
font-weight: inherit !important;
line-height: inherit !important
}
u + #body a {
color: inherit;
text-decoration: none;
font-size: inherit;
font-family: inherit;
font-weight: inherit;
line-height: inherit;
}
a[href^="mailto"],
a[href^="tel"],
a[href^="sms"] {
color: inherit;
text-decoration: none
}
</style>
<style type="text/css">
@media (min-width: 481px) {
.hd { display: none!important }
}
</style>
<style type="text/css">
@media (max-width: 480px) {
.hm { display: none!important }
}
</style>
<style type="text/css">
@media (max-width: 480px) {
.t5{mso-line-height-alt:0px!important;line-height:0!important;display:none!important}.t6{border-top-left-radius:0!important;border-top-right-radius:0!important}.t56{border-bottom-right-radius:0!important;border-bottom-left-radius:0!important}
}
</style>
<!--[if !mso]>-->
<link href="https://fonts.googleapis.com/css2?family=Roboto:wght@400;700&amp;display=swap" rel="stylesheet" type="text/css" />
<!--<![endif]-->
<!--[if mso]>
<xml>
<o:OfficeDocumentSettings>
<o:AllowPNG/>
<o:PixelsPerInch>96</o:PixelsPerInch>
</o:OfficeDocumentSettings>
</xml>
<![endif]-->
</head>
<body id="body" class="t63" style="min-width:100%;Margin:0px;padding:0px;background-color:#292929;"><div class="t62" style="background-color:#292929;"><table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" align="center"><tr><td class="t61" style="font-size:0;line-height:0;mso-line-height-rule:exactly;background-color:#292929;" valign="top" align="center">
<!--[if mso]>
<v:background xmlns:v="urn:schemas-microsoft-com:vml" fill="true" stroke="false">
<v:fill color="#292929"/>
</v:background>
<![endif]-->
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" align="center" id="innerTable"><tr><td><div class="t5" style="mso-line-height-rule:exactly;mso-line-height-alt:100px;line-height:100px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr><tr><td align="center">
<table class="t9" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="600" class="t8" style="width:600px;">
<table class="t7" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t6" style="overflow:hidden;background-color:#D0D7E2;padding:40px 0 40px 0;border-radius:14px 14px 0 0;"><table role="presentation" width="100%" cellpadding="0" cellspacing="0" style="width:100% !important;"><tr><td align="center">
<table class="t4" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="256" class="t3" style="width:256px;">
<table class="t2" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t1"><div style="font-size:0px;"><img class="t0" style="display:block;border:0;height:auto;width:100%;Margin:0;max-width:100%;" width="256" height="67" alt="" src="https://lwe.vote/Images/logo.png"/></div></td></tr></table>
</td></tr></table>
</td></tr></table></td></tr></table>
</td></tr></table>
</td></tr><tr><td align="center">
<table class="t59" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="600" class="t58" style="width:600px;">
<table class="t57" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t56" style="overflow:hidden;background-color:#FFFFFF;padding:40px 30px 40px 30px;border-radius:0 0 14px 14px;"><table role="presentation" width="100%" cellpadding="0" cellspacing="0" style="width:100% !important;"><tr><td align="center">
<table class="t14" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="540" class="t13" style="width:600px;">
<table class="t12" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t11"><p class="t10" style="margin:0;Margin:0;font-family:Roboto,BlinkMacSystemFont,Segoe UI,Helvetica Neue,Arial,sans-serif;line-height:22px;font-weight:400;font-style:normal;font-size:16px;text-decoration:none;text-transform:none;direction:ltr;color:#333333;text-align:left;mso-line-height-rule:exactly;mso-text-raise:2px;">Hello,</p></td></tr></table>
</td></tr></table>
</td></tr><tr><td><div class="t15" style="mso-line-height-rule:exactly;mso-line-height-alt:13px;line-height:13px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr><tr><td align="center">
<table class="t20" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="540" class="t19" style="width:600px;">
<table class="t18" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t17"><p class="t16" style="margin:0;Margin:0;font-family:Roboto,BlinkMacSystemFont,Segoe UI,Helvetica Neue,Arial,sans-serif;line-height:22px;font-weight:400;font-style:normal;font-size:16px;text-decoration:none;text-transform:none;direction:ltr;color:#333333;text-align:left;mso-line-height-rule:exactly;mso-text-raise:2px;">A new legislative committee meeting has been processed and transcribed:</p></td></tr></table>
</td></tr></table>
</td></tr><tr><td><div class="t21" style="mso-line-height-rule:exactly;mso-line-height-alt:11px;line-height:11px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr><tr><td align="center">
<table class="t35" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="540" class="t34" style="width:600px;">
<table class="t33" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t32"><p class="t31" style="margin:0;Margin:0;font-family:Roboto,BlinkMacSystemFont,Segoe UI,Helvetica Neue,Arial,sans-serif;line-height:22px;font-weight:400;font-style:normal;font-size:16px;text-decoration:none;text-transform:none;direction:ltr;color:#333333;text-align:left;mso-line-height-rule:exactly;mso-text-raise:2px;"><span class="t22" style="margin:0;Margin:0;font-weight:700;mso-line-height-rule:exactly;">Committee:</span> {COMMITTEE_NAME} Meeting <br/><span class="t23" style="margin:0;Margin:0;font-weight:700;mso-line-height-rule:exactly;">Date:</span> {MEETING_DATE}<br/><span class="t24" style="margin:0;Margin:0;font-weight:700;mso-line-height-rule:exactly;">Time:</span> {MEETING_TIME} <br/><span class="t25" style="margin:0;Margin:0;font-weight:700;mso-line-height-rule:exactly;">Status:</span> {STATUS} <br/><span class="t26" style="margin:0;Margin:0;font-weight:700;mso-line-height-rule:exactly;">Video Source:</span> <a class="t30" href="{VIDEO_SOURCE_URL}" style="margin:0;Margin:0;font-weight:700;font-style:normal;text-decoration:none;direction:ltr;color:#0000FF;mso-line-height-rule:exactly;" target="_blank">View Meeting Recording</a><br/><span class="t29" style="margin:0;Margin:0;color:#333333;mso-line-height-rule:exactly;"><span class="t27" style="margin:0;Margin:0;font-weight:700;mso-line-height-rule:exactly;">Processed at:</span><span class="t28" style="margin:0;Margin:0;font-weight:400;mso-line-height-rule:exactly;"> {PROCESSING_TIME}</span></span></p></td></tr></table>
</td></tr></table>
</td></tr><tr><td><div class="t36" style="mso-line-height-rule:exactly;mso-line-height-alt:15px;line-height:15px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr><tr><td align="center">
<table class="t43" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="540" class="t42" style="width:600px;">
<table class="t41" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t40"><p class="t39" style="margin:0;Margin:0;font-family:Roboto,BlinkMacSystemFont,Segoe UI,Helvetica Neue,Arial,sans-serif;line-height:22px;font-weight:400;font-style:normal;font-size:16px;text-decoration:none;text-transform:none;direction:ltr;color:#333333;text-align:left;mso-line-height-rule:exactly;mso-text-raise:2px;">The transcript is now available for internal use at the following link: <a class="t38" href="{NEXTCLOUD_LINK}" style="margin:0;Margin:0;font-weight:700;font-style:normal;text-decoration:none;direction:ltr;color:#0000FF;mso-line-height-rule:exactly;" target="_blank"><span class="t37" style="margin:0;Margin:0;mso-line-height-rule:exactly;">{NEXTCLOUD_LINK}</span></a></p></td></tr></table>
</td></tr></table>
</td></tr><tr><td><div class="t44" style="mso-line-height-rule:exactly;mso-line-height-alt:25px;line-height:25px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr><tr><td align="center">
<table class="t49" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="540" class="t48" style="width:600px;">
<table class="t47" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t46"><p class="t45" style="margin:0;Margin:0;font-family:Roboto,BlinkMacSystemFont,Segoe UI,Helvetica Neue,Arial,sans-serif;line-height:22px;font-weight:400;font-style:normal;font-size:16px;text-decoration:none;text-transform:none;direction:ltr;color:#333333;text-align:left;mso-line-height-rule:exactly;mso-text-raise:2px;">Enjoy!</p></td></tr></table>
</td></tr></table>
</td></tr><tr><td><div class="t50" style="mso-line-height-rule:exactly;mso-line-height-alt:6px;line-height:6px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr><tr><td align="center">
<table class="t55" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="540" class="t54" style="width:600px;">
<table class="t53" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t52"><p class="t51" style="margin:0;Margin:0;font-family:Roboto,BlinkMacSystemFont,Segoe UI,Helvetica Neue,Arial,sans-serif;line-height:22px;font-weight:400;font-style:normal;font-size:16px;text-decoration:none;text-transform:none;direction:ltr;color:#333333;text-align:left;mso-line-height-rule:exactly;mso-text-raise:2px;">LWE.Vote</p></td></tr></table>
</td></tr></table>
</td></tr></table></td></tr></table>
</td></tr></table>
</td></tr><tr><td><div class="t60" style="mso-line-height-rule:exactly;mso-line-height-alt:80px;line-height:80px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr></table></td></tr></table></div><div class="gmail-fix" style="display: none; white-space: nowrap; font: 15px courier; line-height: 0;">&nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp;</div></body>
</html>''' class="t29" style="margin:0;Margin:0;color:#333333;mso-line-height-rule:exactly;"><span class="t27" style="margin:0;Margin:0;font-weight:700;mso-line-height-rule:exactly;">Processed at:</span><span class="t28" style="margin:0;Margin:0;font-weight:400;mso-line-height-rule:exactly;"> {PROCESSING_TIME}</span></span></p></td></tr></table>
</td></tr></table>
</td></tr><tr><td><div class="t36" style="mso-line-height-rule:exactly;mso-line-height-alt:15px;line-height:15px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr><tr><td align="center">
<table class="t43" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="540" class="t42" style="width:600px;">
<table class="t41" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t40"><p class="t39" style="margin:0;Margin:0;font-family:Roboto,BlinkMacSystemFont,Segoe UI,Helvetica Neue,Arial,sans-serif;line-height:22px;font-weight:400;font-style:normal;font-size:16px;text-decoration:none;text-transform:none;direction:ltr;color:#333333;text-align:left;mso-line-height-rule:exactly;mso-text-raise:2px;">The transcript is now available for internal use at the following link: <a class="t38" href="{NEXTCLOUD_LINK}" style="margin:0;Margin:0;font-weight:700;font-style:normal;text-decoration:none;direction:ltr;color:#0000FF;mso-line-height-rule:exactly;" target="_blank"><span class="t37" style="margin:0;Margin:0;mso-line-height-rule:exactly;">{NEXTCLOUD_LINK}</span></a></p></td></tr></table>
</td></tr></table>
</td></tr><tr><td><div class="t44" style="mso-line-height-rule:exactly;mso-line-height-alt:25px;line-height:25px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr><tr><td align="center">
<table class="t49" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="540" class="t48" style="width:600px;">
<table class="t47" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t46"><p class="t45" style="margin:0;Margin:0;font-family:Roboto,BlinkMacSystemFont,Segoe UI,Helvetica Neue,Arial,sans-serif;line-height:22px;font-weight:400;font-style:normal;font-size:16px;text-decoration:none;text-transform:none;direction:ltr;color:#333333;text-align:left;mso-line-height-rule:exactly;mso-text-raise:2px;">Enjoy!</p></td></tr></table>
</td></tr></table>
</td></tr><tr><td><div class="t50" style="mso-line-height-rule:exactly;mso-line-height-alt:6px;line-height:6px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr><tr><td align="center">
<table class="t55" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="540" class="t54" style="width:600px;">
<table class="t53" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t52"><p class="t51" style="margin:0;Margin:0;font-family:Roboto,BlinkMacSystemFont,Segoe UI,Helvetica Neue,Arial,sans-serif;line-height:22px;font-weight:400;font-style:normal;font-size:16px;text-decoration:none;text-transform:none;direction:ltr;color:#333333;text-align:left;mso-line-height-rule:exactly;mso-text-raise:2px;">LWE.Vote</p></td></tr></table>
</td></tr></table>
</td></tr></table></td></tr></table>
</td></tr></table>
</td></tr><tr><td><div class="t60" style="mso-line-height-rule:exactly;mso-line-height-alt:80px;line-height:80px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr></table></td></tr></table></div><div class="gmail-fix" style="display: none; white-space: nowrap; font: 15px courier; line-height: 0;">&nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp;</div></body>
</html>"""
                
                # Replace placeholders with actual data
                processing_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S MST')
                
                # Determine status based on transcript result
                status = "Transcription Complete"
                if result.get('share_link'):
                    status += " & Available"
                
                # Get Nextcloud share link
                nextcloud_link = result.get('share_link', '#')
                
                html_content = html_template.replace('{COMMITTEE_NAME}', committee_name)
                html_content = html_content.replace('{MEETING_DATE}', formatted_date)
                html_content = html_content.replace('{MEETING_TIME}', meeting_time)
                html_content = html_content.replace('{STATUS}', status)
                html_content = html_content.replace('{VIDEO_SOURCE_URL}', entry.get('link', '#'))
                html_content = html_content.replace('{PROCESSING_TIME}', processing_time)
                html_content = html_content.replace('{NEXTCLOUD_LINK}', nextcloud_link)
                
                # Create HTML part
                html_part = MIMEText(html_content, 'html')
                msg.attach(html_part)
                
                # Send email
                server = smtplib.SMTP('smtp.gmail.com', 587)
                server.starttls()
                server.login(EMAIL_USER, EMAIL_PASSWORD)
                server.send_message(msg)
                server.quit()
                
                logger.info(f"HTML notification email sent successfully to {recipient} for {committee_name}")
                
    except Exception as e:
        logger.error(f"Error sending HTML notification: {e}")


def test_proxy_connectivity():
    """Test if Oxylabs proxy is working correctly."""
    logger.info("Testing Oxylabs proxy connectivity...")
    
    test_url = "http://httpbin.org/ip"  # Simple IP echo service
    
    if proxy_working and proxy_config:
        proxy_dict = get_proxy_dict(proxy_config)
        
        try:
            response = requests.get(test_url, proxies=proxy_dict, timeout=10)
            response.raise_for_status()
            ip_data = response.json()
            logger.info(f"Oxylabs proxy working - IP: {ip_data.get('origin', 'unknown')}")
        except Exception as e:
            logger.warning(f"Oxylabs proxy test failed: {e}")
    
    # Test without proxy
    try:
        response = requests.get(test_url, timeout=10)
        response.raise_for_status()
        ip_data = response.json()
        logger.info(f"Direct connection - Working (IP: {ip_data.get('origin', 'unknown')})")
    except Exception as e:
        logger.warning(f"Direct connection - Failed: {e}")


def main():
    """Main function to schedule and run the monitoring."""
    logger.info("Starting Legislature page monitoring with yt-dlp, AssemblyAI, Nextcloud integration, and Oxylabs Proxy...")
    
    # Check Nextcloud configuration
    if not all([NEXTCLOUD_URL, NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN]):
        logger.error("Nextcloud configuration missing. Please check NEXTCLOUD_URL, NEXTCLOUD_USERNAME, and NEXTCLOUD_TOKEN in .env file")
        return
    else:
        logger.info(f"Nextcloud configured: {NEXTCLOUD_URL}")
    
    # Initialize Oxylabs proxy system
    logger.info("Initializing Oxylabs proxy system...")
    if not test_oxylabs_proxy():
        logger.warning("Failed to connect to Oxylabs proxy. Script will attempt to run without proxy.")
    else:
        logger.info("Oxylabs proxy system initialized successfully")
        # Test proxy connectivity
        test_proxy_connectivity()
    
    # Check for required dependencies
    try:
        import yt_dlp
        logger.info("yt-dlp is available")
    except ImportError:
        logger.error("yt-dlp not found. Please install with: pip install yt-dlp")
        return
    
    try:
        from docx import Document
        logger.info("python-docx is available")
    except ImportError:
        logger.error("python-docx not found. Please install with: pip install python-docx")
        return
    
    # Check for ffmpeg (optional, for fallback)
    try:
        subprocess.run(['ffmpeg', '-version'], check=True, capture_output=True)
        logger.info("ffmpeg is available (for fallback audio extraction)")
    except (subprocess.SubprocessError, FileNotFoundError):
        logger.warning("ffmpeg not found. Audio extraction will rely on yt-dlp only.")
    
    if not ASSEMBLYAI_API_KEY:
        logger.warning("AssemblyAI API key not found. Transcription will not work.")
    
    logger.info("All dependencies checked. Starting monitoring system...")
    
    # Run immediately on startup
    check_for_new_entries()
    
    # Schedule checks every few hours, Monday through Friday only (GMT times for MST = GMT-7)
    gmt_slots = ["16:00", "18:00", "20:00", "22:00", "00:00", "02:00"]

    for t in gmt_slots:
        schedule.every().monday.at(t).do(check_for_new_entries)
        schedule.every().tuesday.at(t).do(check_for_new_entries)
        schedule.every().wednesday.at(t).do(check_for_new_entries)
        schedule.every().thursday.at(t).do(check_for_new_entries)
        schedule.every().friday.at(t).do(check_for_new_entries)
    
    logger.info(f"Scheduled monitoring Monday-Friday at GMT times: {', '.join(gmt_slots)}")
    logger.info("Legislature monitoring system is now running (weekdays only)...")
    
    # Keep the script running
    while True:
        schedule.run_pending()
        time.sleep(60)


if __name__ == "__main__":
    main()
                
                # Create HTML part
                html_part = MIMEText(html_content, 'html')
                msg.attach(html_part)
                
                # Send email
                server = smtplib.SMTP('smtp.gmail.com', 587)
                server.starttls()
                server.login(EMAIL_USER, EMAIL_PASSWORD)
                server.send_message(msg)
                server.quit()
                
                logger.info(f"HTML notification email sent successfully to {recipient} for {committee_name}")
                
    except Exception as e:
        logger.error(f"Error sending HTML notification: {e}")="0" border="0" align="center" id="innerTable"><tr><td><div class="t5" style="mso-line-height-rule:exactly;mso-line-height-alt:100px;line-height:100px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr><tr><td align="center">
<table class="t9" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="600" class="t8" style="width:600px;">
<table class="t7" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t6" style="overflow:hidden;background-color:#D0D7E2;padding:40px 0 40px 0;border-radius:14px 14px 0 0;"><table role="presentation" width="100%" cellpadding="0" cellspacing="0" style="width:100% !important;"><tr><td align="center">
<table class="t4" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="256" class="t3" style="width:256px;">
<table class="t2" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t1"><div style="font-size:0px;"><img class="t0" style="display:block;border:0;height:auto;width:100%;Margin:0;max-width:100%;" width="256" height="67" alt="" src="https://lwe.vote/Images/logo.png"/></div></td></tr></table>
</td></tr></table>
</td></tr></table></td></tr></table>
</td></tr></table>
</td></tr><tr><td align="center">
<table class="t59" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="600" class="t58" style="width:600px;">
<table class="t57" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t56" style="overflow:hidden;background-color:#FFFFFF;padding:40px 30px 40px 30px;border-radius:0 0 14px 14px;"><table role="presentation" width="100%" cellpadding="0" cellspacing="0" style="width:100% !important;"><tr><td align="center">
<table class="t14" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="540" class="t13" style="width:600px;">
<table class="t12" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t11"><p class="t10" style="margin:0;Margin:0;font-family:Roboto,BlinkMacSystemFont,Segoe UI,Helvetica Neue,Arial,sans-serif;line-height:22px;font-weight:400;font-style:normal;font-size:16px;text-decoration:none;text-transform:none;direction:ltr;color:#333333;text-align:left;mso-line-height-rule:exactly;mso-text-raise:2px;">Hello,</p></td></tr></table>
</td></tr></table>
</td></tr><tr><td><div class="t15" style="mso-line-height-rule:exactly;mso-line-height-alt:13px;line-height:13px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr><tr><td align="center">
<table class="t20" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="540" class="t19" style="width:600px;">
<table class="t18" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t17"><p class="t16" style="margin:0;Margin:0;font-family:Roboto,BlinkMacSystemFont,Segoe UI,Helvetica Neue,Arial,sans-serif;line-height:22px;font-weight:400;font-style:normal;font-size:16px;text-decoration:none;text-transform:none;direction:ltr;color:#333333;text-align:left;mso-line-height-rule:exactly;mso-text-raise:2px;">A new legislative committee meeting has been processed and transcribed:</p></td></tr></table>
</td></tr></table>
</td></tr><tr><td><div class="t21" style="mso-line-height-rule:exactly;mso-line-height-alt:11px;line-height:11px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr><tr><td align="center">
<table class="t35" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="540" class="t34" style="width:600px;">
<table class="t33" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t32"><p class="t31" style="margin:0;Margin:0;font-family:Roboto,BlinkMacSystemFont,Segoe UI,Helvetica Neue,Arial,sans-serif;line-height:22px;font-weight:400;font-style:normal;font-size:16px;text-decoration:none;text-transform:none;direction:ltr;color:#333333;text-align:left;mso-line-height-rule:exactly;mso-text-raise:2px;"><span class="t22" style="margin:0;Margin:0;font-weight:700;mso-line-height-rule:exactly;">Committee:</span> {COMMITTEE_NAME} Meeting <br/><span class="t23" style="margin:0;Margin:0;font-weight:700;mso-line-height-rule:exactly;">Date:</span> {MEETING_DATE}<br/><span class="t24" style="margin:0;Margin:0;font-weight:700;mso-line-height-rule:exactly;">Time:</span> {MEETING_TIME} <br/><span class="t25" style="margin:0;Margin:0;font-weight:700;mso-line-height-rule:exactly;">Status:</span> {STATUS} <br/><span class="t26" style="margin:0;Margin:0;font-weight:700;mso-line-height-rule:exactly;">Video Source:</span> <a class="t30" href="{VIDEO_SOURCE_URL}" style="margin:0;Margin:0;font-weight:700;font-style:normal;text-decoration:none;direction:ltr;color:#0000FF;mso-line-height-rule:exactly;" target="_blank">View Meeting Recording</a><br/><span class="t29" style="margin:0;Margin:0;color:#333333;mso-line-height-rule:exactly;"><span class="t27" style="margin:0;Margin:0;font-weight:700;mso-line-height-rule:exactly;">Processed at:</span><span class="t28" style="margin:0;Margin:0;font-weight:400;mso-line-height-rule:exactly;"> {PROCESSING_TIME}</span></span></p></td></tr></table>
</td></tr></table>
</td></tr><tr><td><div class="t36" style="mso-line-height-rule:exactly;mso-line-height-alt:15px;line-height:15px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr><tr><td align="center">
<table class="t43" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="540" class="t42" style="width:600px;">
<table class="t41" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t40"><p class="t39" style="margin:0;Margin:0;font-family:Roboto,BlinkMacSystemFont,Segoe UI,Helvetica Neue,Arial,sans-serif;line-height:22px;font-weight:400;font-style:normal;font-size:16px;text-decoration:none;text-transform:none;direction:ltr;color:#333333;text-align:left;mso-line-height-rule:exactly;mso-text-raise:2px;">The transcript is now available for internal use at the following link: <a class="t38" href="{NEXTCLOUD_LINK}" style="margin:0;Margin:0;font-weight:700;font-style:normal;text-decoration:none;direction:ltr;color:#0000FF;mso-line-height-rule:exactly;" target="_blank"><span class="t37" style="margin:0;Margin:0;mso-line-height-rule:exactly;">Access Transcript</span></a></p></td></tr></table>
</td></tr></table>
</td></tr><tr><td><div class="t44" style="mso-line-height-rule:exactly;mso-line-height-alt:25px;line-height:25px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr><tr><td align="center">
<table class="t49" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="540" class="t48" style="width:600px;">
<table class="t47" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t46"><p class="t45" style="margin:0;Margin:0;font-family:Roboto,BlinkMacSystemFont,Segoe UI,Helvetica Neue,Arial,sans-serif;line-height:22px;font-weight:400;font-style:normal;font-size:16px;text-decoration:none;text-transform:none;direction:ltr;color:#333333;text-align:left;mso-line-height-rule:exactly;mso-text-raise:2px;">Enjoy!</p></td></tr></table>
</td></tr></table>
</td></tr><tr><td><div class="t50" style="mso-line-height-rule:exactly;mso-line-height-alt:6px;line-height:6px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr><tr><td align="center">
<table class="t55" role="presentation" cellpadding="0" cellspacing="0" style="Margin-left:auto;Margin-right:auto;"><tr><td width="540" class="t54" style="width:600px;">
<table class="t53" role="presentation" cellpadding="0" cellspacing="0" width="100%" style="width:100%;"><tr><td class="t52"><p class="t51" style="margin:0;Margin:0;font-family:Roboto,BlinkMacSystemFont,Segoe UI,Helvetica Neue,Arial,sans-serif;line-height:22px;font-weight:400;font-style:normal;font-size:16px;text-decoration:none;text-transform:none;direction:ltr;color:#333333;text-align:left;mso-line-height-rule:exactly;mso-text-raise:2px;">LWE.Vote</p></td></tr></table>
</td></tr></table>
</td></tr></table></td></tr></table>
</td></tr></table>
</td></tr><tr><td><div class="t60" style="mso-line-height-rule:exactly;mso-line-height-alt:80px;line-height:80px;font-size:1px;display:block;">&nbsp;&nbsp;</div></td></tr></table></td></tr></table></div><div class="gmail-fix" style="display: none; white-space: nowrap; font: 15px courier; line-height: 0;">&nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp;</div></body>
</html>"""
                
                # Replace placeholders with actual data
                processing_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S MST')
                
                # Determine status based on transcript result
                status = "Transcription Complete"
                if result.get('share_link'):
                    status += " & Available"
                
                html_content = html_template.replace('{COMMITTEE_NAME}', committee_name)
                html_content = html_content.replace('{MEETING_DATE}', formatted_date)
                html_content = html_content.replace('{MEETING_TIME}', meeting_time)
                html_content = html_content.replace('{STATUS}', status)
                html_content = html_content.replace('{VIDEO_SOURCE_URL}', entry.get('link', '#'))
                html_content = html_content.replace('{PROCESSING_TIME}', processing_time)
                html_content = html_content.replace('{NEXTCLOUD_LINK}', result.get('share_link', '#'))
                
                # Create HTML part
                html_part = MIMEText(html_content, 'html')
                msg.attach(html_part)
                
                # Send email
                server = smtplib.SMTP('smtp.gmail.com', 587)
                server.starttls()
                server.login(EMAIL_USER, EMAIL_PASSWORD)
                server.send_message(msg)
                server.quit()
                
                logger.info(f"HTML notification email sent successfully to {recipient} for {committee_name}")
                
    except Exception as e:
        logger.error(f"Error sending HTML notification: {e}")


def extract_hls_stream_url(entry_url):
    """Extract the HLS stream URL from the entry page using proxy."""
    try:
        logger.info(f"Accessing entry page to extract HLS stream: {entry_url}")
        
        # Use proxy-enabled request
        response = make_request_with_proxy(entry_url)

        # First try to find the availableStreams JSON data
        available_streams_pattern = r'var\s+availableStreams\s*=\s*(\[.*?\]);'
        available_streams_match = re.search(available_streams_pattern,
                                            response.text,
                                            re.DOTALL)
        if available_streams_match:
            streams_json_str = available_streams_match.group(1)
            try:
                streams_data = json.loads(streams_json_str)
                logger.info(f"Found {len(streams_data)} streams in availableStreams JSON")
                
                # Look for the best stream (non-live, enabled, with URL ending in .m3u8)
                best_stream = None
                for stream in streams_data:
                    url = stream.get('Url', '')
                    is_live = stream.get('IsLive', True)
                    enabled = stream.get('Enabled', False)
                    
                    logger.info(f"Stream: URL={url[:80]}{'...' if len(url) > 80 else ''}, IsLive={is_live}, Enabled={enabled}")
                    
                    if url.endswith('.m3u8') and enabled and not is_live:
                        # This is a good recorded stream
                        best_stream = stream
                        break
                    elif url.endswith('.m3u8') and enabled:
                        # Fallback to any enabled stream
                        if best_stream is None:
                            best_stream = stream
                
                if best_stream:
                    # Clean up the URL - replace escaped forward slashes but preserve URL encoding
                    hls_url = best_stream['Url'].replace('\\/', '/')
                    logger.info(f"Selected stream: IsLive={best_stream.get('IsLive')}, Duration={best_stream.get('Duration')}s")
                    logger.info(f"Extracted HLS URL from JSON: {hls_url}")
                    return hls_url
                else:
                    logger.warning("No suitable streams found in availableStreams JSON")
                    
            except json.JSONDecodeError as e:
                logger.error(f"Could not parse availableStreams JSON: {e}")
                logger.debug(f"JSON string was: {streams_json_str[:200]}...")

        # Enhanced fallback patterns - look for VOD URLs first
        vod_patterns = [
            # Pattern for VOD URLs with _definst_ and encoded title (most specific)
            r'(https?://[^"\'\s]+vod-2/_definst_/[^"\'\s]+\.mp4/playlist\.m3u8)',
            # Pattern for any VOD URL
            r'(https?://[^"\'\s]+vod[^"\'\s]*\.m3u8)',
            # General m3u8 pattern but prioritize those with dates/times
            r'(https?://[^"\'\s]+/\d{4}/\d{2}/\d{2}/[^"\'\s]+\.m3u8)',
        ]
        
        for i, pattern in enumerate(vod_patterns):
            matches = re.findall(pattern, response.text, re.IGNORECASE)
            if matches:
                logger.info(f"Found {len(matches)} matches with pattern {i+1}")
                # Filter out live streams if we have multiple matches
                for match in matches:
                    if 'live' not in match.lower():
                        logger.info(f"Found VOD HLS URL via regex pattern {i+1}: {match}")
                        return match
                # If all have 'live', take the first one
                logger.info(f"Found HLS URL via regex pattern {i+1} (might be live): {matches[0]}")
                return matches[0]

        # Last resort - any m3u8 URL
        general_m3u8_pattern = r'(https?://[^"\'\s]+\.m3u8)'
        matches = re.findall(general_m3u8_pattern, response.text)
        if matches:
            logger.info(f"Found {len(matches)} m3u8 URLs as last resort")
            # Filter out known live stream patterns
            for match in matches:
                if '/live/' not in match and 'playlist.m3u8' in match:
                    logger.info(f"Found potential HLS URL via fallback: {match}")
                    return match
            
            # If no good matches, return first one
            logger.warning(f"Using first available HLS URL (may be live stream): {matches[0]}")
            return matches[0]

        logger.error("No HLS stream URLs found on the page")
        return None

    except Exception as e:
        logger.error(f"Error extracting HLS stream URL: {e}")
        return None


def download_hls_with_ytdlp(hls_url, output_file, max_retries=3):
    """Download HLS stream using yt-dlp with Oxylabs proxy."""
    try:
        output_dir = os.path.dirname(output_file)
        if output_dir:
            Path(output_dir).mkdir(exist_ok=True)
        
        logger.info(f"Downloading HLS stream with yt-dlp from {hls_url} to {output_file}")
        
        # Try multiple attempts
        for attempt in range(max_retries):
            try:
                # Configure yt-dlp options
                ydl_opts = {
                    'outtmpl': output_file.replace('.mp4', '.%(ext)s'),
                    'format': 'best[ext=mp4]/best',  # Prefer mp4, fallback to best available
                    'writesubtitles': False,
                    'writeautomaticsub': False,
                    'ignoreerrors': False,
                    'no_warnings': False,
                    'extractaudio': False,
                    'socket_timeout': 60,
                    'retries': 3,
                    'fragment_retries': 5,
                    'http_chunk_size': 10485760,  # 10MB chunks
                    'hls_prefer_native': True,    # Use native HLS extractor
                }
                
                # Add proxy configuration if available
                if proxy_working and proxy_config:
                    proxy_url = get_proxy_url(proxy_config)
                    ydl_opts['proxy'] = proxy_url
                    logger.info(f"Attempt {attempt + 1}: Using Oxylabs proxy {proxy_config['ip']}:{proxy_config['port']} with yt-dlp")
                else:
                    logger.warning(f"Attempt {attempt + 1}: Proxy not available, using direct connection")
                
                # Add user agent
                ydl_opts['http_headers'] = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                
                # Download with yt-dlp
                with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                    ydl.download([hls_url])
                
                # Check if file was created (yt-dlp might change extension)
                possible_files = [
                    output_file,
                    output_file.replace('.mp4', '.mkv'),
                    output_file.replace('.mp4', '.webm'),
                    output_file.replace('.mp4', '.m4v')
                ]
                
                downloaded_file = None
                for possible_file in possible_files:
                    if os.path.exists(possible_file) and os.path.getsize(possible_file) > 1024:
                        downloaded_file = possible_file
                        break
                
                if downloaded_file:
                    # Rename to desired output file if necessary
                    if downloaded_file != output_file:
                        os.rename(downloaded_file, output_file)
                        downloaded_file = output_file
                    
                    file_size_mb = os.path.getsize(downloaded_file) / (1024 * 1024)
                    logger.info(f"Successfully downloaded video to {downloaded_file} (size: {file_size_mb:.2f} MB)")
                    return downloaded_file
                else:
                    logger.error("Download failed - no output file found")
                    
            except Exception as e:
                logger.warning(f"yt-dlp attempt {attempt + 1} failed: {e}")
                if attempt < max_retries - 1:
                    wait_time = (2 ** attempt) + random.uniform(1, 3)
                    logger.info(f"Waiting {wait_time:.1f} seconds before retry...")
                    time.sleep(wait_time)
                continue
        
        logger.error(f"All {max_retries} yt-dlp attempts failed")
        return None
        
    except Exception as e:
        logger.error(f"Error in yt-dlp download process: {e}")
        return None


def download_with_ytdlp_fallback(hls_url, output_file):
    """Download with yt-dlp, fallback to direct if proxy fails."""
    try:
        # First try with proxy
        result = download_hls_with_ytdlp(hls_url, output_file)
        if result:
            return result
        
        # Fallback to direct connection
        logger.info("Trying yt-dlp with direct connection...")
        
        ydl_opts = {
            'outtmpl': output_file.replace('.mp4', '.%(ext)s'),
            'format': 'best[ext=mp4]/best',
            'writesubtitles': False,
            'writeautomaticsub': False,
            'ignoreerrors': False,
            'socket_timeout': 60,
            'retries': 5,
            'fragment_retries': 10,
            'http_chunk_size': 10485760,
            'hls_prefer_native': True,
            'http_headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([hls_url])
        
        # Check for downloaded file
        possible_files = [
            output_file,
            output_file.replace('.mp4', '.mkv'),
            output_file.replace('.mp4', '.webm'),
            output_file.replace('.mp4', '.m4v')
        ]
        
        for possible_file in possible_files:
            if os.path.exists(possible_file) and os.path.getsize(possible_file) > 1024:
                if possible_file != output_file:
                    os.rename(possible_file, output_file)
                
                file_size_mb = os.path.getsize(output_file) / (1024 * 1024)
                logger.info(f"Direct connection successful (size: {file_size_mb:.2f} MB)")
                return output_file
        
        logger.error("Direct connection also failed")
        return None
        
    except Exception as e:
        logger.error(f"yt-dlp fallback failed: {e}")
        return None


def download_video(entry_url):
    """Download video from the entry page using HLS stream extraction with yt-dlp."""
    try:
        hls_url = extract_hls_stream_url(entry_url)
        if not hls_url:
            logger.error("Could not find HLS stream URL on the page")
            return None
        
        video_filename = os.path.join(
            DOWNLOAD_DIR, 
            f"video_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.mp4"
        )
        
        # Use yt-dlp with fallback strategies
        return download_with_ytdlp_fallback(hls_url, video_filename)
    
    except Exception as e:
        logger.error(f"Error in video download process: {e}")
        return None


def extract_audio_from_video(video_path):
    """Extract audio from video file using ffmpeg directly (more reliable for local files)."""
    try:
        if not os.path.exists(video_path):
            logger.error(f"Video file not found: {video_path}")
            return None
        
        audio_path = video_path.replace('.mp4', '.mp3')
        
        logger.info(f"Extracting audio from {video_path} to {audio_path}")
        
        # Use ffmpeg directly for local file conversion (more reliable than yt-dlp for this)
        try:
            result = subprocess.run([
                'ffmpeg', '-i', video_path, 
                '-vn',  # No video
                '-acodec', 'libmp3lame',  # MP3 codec
                '-ab', '192k',  # 192kbps bitrate
                '-ar', '44100',  # Sample rate
                '-y',  # Overwrite output file
                audio_path
            ], check=True, capture_output=True, text=True)
            
            if os.path.exists(audio_path) and os.path.getsize(audio_path) > 1024:
                file_size_mb = os.path.getsize(audio_path) / (1024 * 1024)
                logger.info(f"Audio extracted successfully to {audio_path} (size: {file_size_mb:.2f} MB)")
                return audio_path
            else:
                logger.error("Audio extraction failed - output file not created or too small")
                return None
                
        except subprocess.CalledProcessError as e:
            logger.error(f"ffmpeg failed: {e.stderr}")
            # Try alternative ffmpeg method
            return extract_audio_with_alternative_ffmpeg(video_path)
            
        except FileNotFoundError:
            logger.error("ffmpeg not found. Cannot extract audio.")
            return None
    
    except Exception as e:
        logger.error(f"Error extracting audio: {e}")
        return None


def extract_audio_with_alternative_ffmpeg(video_path):
    """Alternative ffmpeg method with different parameters."""
    try:
        audio_path = video_path.replace('.mp4', '.mp3')
        
        logger.info(f"Trying alternative ffmpeg method...")
        
        # Alternative method with simpler parameters
        result = subprocess.run([
            'ffmpeg', '-i', video_path, 
            '-q:a', '2',  # High quality audio
            '-map', 'a',  # Map audio stream
            '-y',  # Overwrite
            audio_path
        ], check=True, capture_output=True, text=True)
        
        if os.path.exists(audio_path) and os.path.getsize(audio_path) > 1024:
            file_size_mb = os.path.getsize(audio_path) / (1024 * 1024)
            logger.info(f"Audio extracted with alternative method to {audio_path} (size: {file_size_mb:.2f} MB)")
            return audio_path
        else:
            logger.error("Alternative audio extraction failed")
            return None
            
    except Exception as e:
        logger.error(f"Alternative ffmpeg method failed: {e}")
        return None


def transcribe_with_assemblyai(audio_path):
    """Transcribe audio using AssemblyAI API."""
    if not ASSEMBLYAI_API_KEY:
        logger.error("AssemblyAI API key not found in environment variables")
        return None
    
    try:
        headers = {
            "authorization": ASSEMBLYAI_API_KEY,
            "content-type": "application/json"
        }
        
        logger.info(f"Uploading audio file to AssemblyAI: {audio_path}")
        
        with open(audio_path, "rb") as f:
            response = requests.post(
                f"{ASSEMBLYAI_BASE_URL}/upload",
                headers=headers,
                data=f
            )
            response.raise_for_status()
        
        upload_url = response.json()["upload_url"]
        logger.info(f"Audio uploaded successfully. URL: {upload_url}")
        
        logger.info("Submitting transcription request")
        response = requests.post(
            f"{ASSEMBLYAI_BASE_URL}/transcript",
            headers=headers,
            json={
                "audio_url": upload_url,
                "speaker_labels": True,
                "punctuate": True,
                "format_text": True
            }
        )
        response.raise_for_status()
        
        transcript_id = response.json()["id"]
        logger.info(f"Transcription request submitted. ID: {transcript_id}")
        
        logger.info("Waiting for transcription to complete...")
        
        while True:
            response = requests.get(
                f"{ASSEMBLYAI_BASE_URL}/transcript/{transcript_id}",
                headers=headers
            )
            response.raise_for_status()
            
            status = response.json()["status"]
            
            if status == "completed":
                logger.info("Transcription completed successfully")
                return response.json()
            elif status == "error":
                logger.error(f"Transcription failed: {response.json().get('error')}")
                return None
            
            logger.info(f"Transcription status: {status}. Waiting...")
            time.sleep(30)  # Poll every 30 seconds
    
    except Exception as e:
        logger.error(f"Error in transcription process: {e}")
        return None


def process_new_entry(entry):
    """Process a new entry: download video, transcribe, and save to Nextcloud."""
    try:
        if not entry.get('link'):
            logger.error("Entry has no link, cannot process")
            return None
        
        video_path = download_video(entry['link'])
        if not video_path:
            logger.error("Failed to download video")
            return None
        
        audio_path = extract_audio_from_video(video_path)
        if not audio_path:
            logger.error("Failed to extract audio")
            return None
        
        transcript_data = transcribe_with_assemblyai(audio_path)
        if not transcript_data:
            logger.error("Failed to transcribe audio")
            return None
        
        # Clean title for use in filename and document
        title_text = entry['text']
        clean_title = re.sub(r'\s+', ' ', title_text).strip()
        
        nextcloud_result = save_transcript_to_nextcloud(transcript_data, clean_title, entry['link'])
        
        entry['transcription'] = {
            'video_path': video_path,
            'audio_path': audio_path,
            'nextcloud_result': nextcloud_result,
            'processed_at': datetime.datetime.now().isoformat()
        }
        
        write_processed_entry(entry)
        
        return nextcloud_result
    
    except Exception as e:
        logger.error(f"Error processing entry: {e}")
        return None


def test_proxy_connectivity():
    """Test if Oxylabs proxy is working correctly."""
    logger.info("Testing Oxylabs proxy connectivity...")
    
    test_url = "http://httpbin.org/ip"  # Simple IP echo service
    
    if proxy_working and proxy_config:
        proxy_dict = get_proxy_dict(proxy_config)
        
        try:
            response = requests.get(test_url, proxies=proxy_dict, timeout=10)
            response.raise_for_status()
            ip_data = response.json()
            logger.info(f"Oxylabs proxy working - IP: {ip_data.get('origin', 'unknown')}")
        except Exception as e:
            logger.warning(f"Oxylabs proxy test failed: {e}")
    
    # Test without proxy
    try:
        response = requests.get(test_url, timeout=10)
        response.raise_for_status()
        ip_data = response.json()
        logger.info(f"Direct connection - Working (IP: {ip_data.get('origin', 'unknown')})")
    except Exception as e:
        logger.warning(f"Direct connection - Failed: {e}")


def is_test_meeting(entry_text):
    """Check if an entry is a test meeting that should be skipped."""
    # Convert to lowercase for case-insensitive matching
    text_lower = entry_text.lower()
    
    # Check for "test meeting" anywhere in the text
    if "test meeting" in text_lower:
        return True
    
    return False


def check_for_new_entries():
    """Check for new entries and process them if found."""
    # Check if we need to test proxy
    if should_update_proxy_list():
        logger.info("Testing Oxylabs proxy...")
        test_oxylabs_proxy()
    
    # Check if we need to run daily cleanup
    if should_run_daily_cleanup():
        logger.info("Running daily cleanup of downloads directory...")
        cleanup_downloads_directory()
    
    logger.info(f"Checking for new entries at {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    current_entries = get_current_entries()
    stored_entries = read_stored_entries()
    
    stored_texts = [entry['text'] for entry in stored_entries]
    new_entries = [entry for entry in current_entries if entry['text'] not in stored_texts]
    
    # Filter out test meetings
    filtered_new_entries = []
    skipped_test_meetings = 0
    
    for entry in new_entries:
        if is_test_meeting(entry['text']):
            logger.info(f"Skipping test meeting: {entry['text'][:50]}...")
            skipped_test_meetings += 1
        else:
            filtered_new_entries.append(entry)
    
    if skipped_test_meetings > 0:
        logger.info(f"Skipped {skipped_test_meetings} test meeting(s)")
    
    if filtered_new_entries:
        logger.info(f"Found {len(filtered_new_entries)} new entries to process!")
        
        transcript_results = []
        for entry in filtered_new_entries:
            logger.info(f"Processing entry: {entry['text'][:50]}...")
            result = process_new_entry(entry)
            transcript_results.append(result)
        
        send_notification(filtered_new_entries, transcript_results)
        write_entries(current_entries)  # Still save all entries (including test meetings) to avoid reprocessing
    else:
        if new_entries:
            logger.info("All new entries were test meetings - no processing needed")
        else:
            logger.info("No new entries found")


def main():
    """Main function to schedule and run the monitoring."""
    logger.info("Starting Legislature page monitoring with yt-dlp, AssemblyAI, Nextcloud integration, and Oxylabs Proxy...")
    
    # Check Nextcloud configuration
    if not all([NEXTCLOUD_URL, NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN]):
        logger.error("Nextcloud configuration missing. Please check NEXTCLOUD_URL, NEXTCLOUD_USERNAME, and NEXTCLOUD_TOKEN in .env file")
        return
    else:
        logger.info(f"Nextcloud configured: {NEXTCLOUD_URL}")
    
    # Initialize Oxylabs proxy system
    logger.info("Initializing Oxylabs proxy system...")
    if not test_oxylabs_proxy():
        logger.warning("Failed to connect to Oxylabs proxy. Script will attempt to run without proxy.")
    else:
        logger.info("Oxylabs proxy system initialized successfully")
        # Test proxy connectivity
        test_proxy_connectivity()
    
    # Check for required dependencies
    try:
        import yt_dlp
        logger.info("yt-dlp is available")
    except ImportError:
        logger.error("yt-dlp not found. Please install with: pip install yt-dlp")
        return
    
    try:
        from docx import Document
        logger.info("python-docx is available")
    except ImportError:
        logger.error("python-docx not found. Please install with: pip install python-docx")
        return
    
    # Check for ffmpeg (optional, for fallback)
    try:
        subprocess.run(['ffmpeg', '-version'], check=True, capture_output=True)
        logger.info("ffmpeg is available (for fallback audio extraction)")
    except (subprocess.SubprocessError, FileNotFoundError):
        logger.warning("ffmpeg not found. Audio extraction will rely on yt-dlp only.")
    
    if not ASSEMBLYAI_API_KEY:
        logger.warning("AssemblyAI API key not found. Transcription will not work.")
    
    logger.info("All dependencies checked. Starting monitoring system...")
    
    # Run immediately on startup
    check_for_new_entries()
    
    # Schedule checks every few hours, Monday through Friday only (GMT times for MST = GMT-7)
    gmt_slots = ["16:00", "18:00", "20:00", "22:00", "00:00", "02:00"]

    for t in gmt_slots:
        schedule.every().monday.at(t).do(check_for_new_entries)
        schedule.every().tuesday.at(t).do(check_for_new_entries)
        schedule.every().wednesday.at(t).do(check_for_new_entries)
        schedule.every().thursday.at(t).do(check_for_new_entries)
        schedule.every().friday.at(t).do(check_for_new_entries)
    
    logger.info(f"Scheduled monitoring Monday-Friday at GMT times: {', '.join(gmt_slots)}")
    logger.info("Legislature monitoring system is now running (weekdays only)...")
    
    # Keep the script running
    while True:
        schedule.run_pending()
        time.sleep(60)


if __name__ == "__main__":
    main()
