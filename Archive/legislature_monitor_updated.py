#!/usr/bin/env python3.12
"""
New Mexico Legislature Monitoring Script with yt-dlp, Oxylabs Proxy, and Nextcloud Integration
Updated to use Nextcloud instead of Google Drive/Docs
"""

import requests
from bs4 import BeautifulSoup
import time
import schedule
import os
import datetime
import smtplib
import json
import re
import urllib.parse
import subprocess
import glob
import random
import yt_dlp
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from dotenv import load_dotenv
import logging
from pathlib import Path
import tempfile
from io import BytesIO

# Document creation imports
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("legislature_monitor.log"),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger(__name__)

# Load environment variables from .env file
load_dotenv()

# URL to monitor
URL = "https://sg001-harmony.sliq.net/00293/Harmony/en/View/RecentEnded/20250522/-1"

# File to store the latest entries
ENTRIES_FILE = "latest_entries.txt"
PROCESSED_ENTRIES_FILE = "processed_entries.txt"
DOWNLOAD_DIR = "downloads"
LAST_CLEANUP_FILE = "last_cleanup.txt"
PROXY_LIST_FILE = "proxy_list.txt"
LAST_PROXY_UPDATE_FILE = "last_proxy_update.txt"

# Email configuration
EMAIL_USER = os.getenv("EMAIL_USER")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")
EMAIL_RECIPIENTS = ["joshua@lwe.digital", "skyedevore@gmail.com"]

# AssemblyAI API configuration
ASSEMBLYAI_API_KEY = os.getenv("ASSEMBLYAI_API_KEY")
ASSEMBLYAI_BASE_URL = "https://api.assemblyai.com/v2"

# Oxylabs Proxy configuration
OXYLABS_PROXY_HOST = "dc.oxylabs.io"
OXYLABS_PROXY_PORT = "8000"
OXYLABS_USERNAME = os.getenv("OXYLABS_USERNAME", "user-dapengi_IAs3C-country-US")
OXYLABS_PASSWORD = os.getenv("OXYLABS_PASSWORD", "XDVLR8SYrj2Ee7+")
OXYLABS_LOCATION_URL = "https://ip.oxylabs.io/location"
PROXY_UPDATE_INTERVAL_HOURS = 6  # Check proxy status every 6 hours

# Nextcloud configuration
NEXTCLOUD_URL = os.getenv("NEXTCLOUD_URL")
NEXTCLOUD_USERNAME = os.getenv("NEXTCLOUD_USERNAME")
NEXTCLOUD_TOKEN = os.getenv("NEXTCLOUD_TOKEN")
NEXTCLOUD_BASE_FOLDER = "Legislative Transcription"

# Create download directory if it doesn't exist
Path(DOWNLOAD_DIR).mkdir(exist_ok=True)

# Global proxy configuration
proxy_config = None
proxy_working = False


def test_oxylabs_proxy():
    """Test if Oxylabs proxy is working and get IP location info."""
    global proxy_config, proxy_working
    
    try:
        logger.info("Testing Oxylabs proxy connection...")
        
        # Create proxy configuration using working method (Method 2 from test)
        proxy_config = {
            'ip': OXYLABS_PROXY_HOST,
            'port': OXYLABS_PROXY_PORT,
            'username': OXYLABS_USERNAME,
            'password': OXYLABS_PASSWORD
        }
        
        # Test the proxy using the working format from our test
        proxy_url = f"http://{OXYLABS_USERNAME}:{OXYLABS_PASSWORD}@{OXYLABS_PROXY_HOST}:{OXYLABS_PROXY_PORT}"
        proxies = {
            "http": proxy_url,
            "https": proxy_url
        }
        
        response = requests.get(
            OXYLABS_LOCATION_URL,
            proxies=proxies,
            timeout=30
        )
        response.raise_for_status()
        
        location_data = response.json()
        
        # Parse the complex Oxylabs response format
        proxy_ip = location_data.get('ip', 'unknown')
        
        # Extract country/city from the providers (try multiple sources)
        country = 'unknown'
        city = 'unknown'
        
        providers = location_data.get('providers', {})
        for provider_name, provider_data in providers.items():
            if provider_data.get('country') and country == 'unknown':
                country = provider_data.get('country', 'unknown')
            if provider_data.get('city') and city == 'unknown':
                city = provider_data.get('city', 'unknown')
        
        logger.info(f"Oxylabs proxy working - IP: {proxy_ip}, Country: {country}, City: {city}")
        
        proxy_working = True
        
        # Save proxy test result
        with open(LAST_PROXY_UPDATE_FILE, 'w') as f:
            f.write(datetime.datetime.now().isoformat())
        
        return True
        
    except Exception as e:
        logger.error(f"Oxylabs proxy test failed: {e}")
        proxy_working = False
        return False


def should_update_proxy_list():
    """Check if we should test the proxy again."""
    try:
        if not os.path.exists(LAST_PROXY_UPDATE_FILE):
            return True
        
        with open(LAST_PROXY_UPDATE_FILE, 'r') as f:
            last_update_str = f.read().strip()
        
        last_update = datetime.datetime.fromisoformat(last_update_str)
        now = datetime.datetime.now()
        
        hours_since_update = (now - last_update).total_seconds() / 3600
        return hours_since_update >= PROXY_UPDATE_INTERVAL_HOURS
        
    except Exception as e:
        logger.error(f"Error checking proxy update time: {e}")
        return True


def get_proxy_dict(proxy):
    """Convert proxy info to requests-compatible proxy dict using working Oxylabs format."""
    if not proxy:
        return None
    
    if proxy['username'] and proxy['password']:
        # Working format: http://username:password@host:port (Method 2 from test)
        proxy_url = f"http://{proxy['username']}:{proxy['password']}@{proxy['ip']}:{proxy['port']}"
    else:
        # Without authentication
        proxy_url = f"http://{proxy['ip']}:{proxy['port']}"
    
    return {
        'http': proxy_url,
        'https': proxy_url
    }


def get_proxy_url(proxy):
    """Convert proxy info to URL format for yt-dlp using working Oxylabs format."""
    if not proxy:
        return None
    
    if proxy['username'] and proxy['password']:
        return f"http://{proxy['username']}:{proxy['password']}@{proxy['ip']}:{proxy['port']}"
    else:
        return f"http://{proxy['ip']}:{proxy['port']}"


def make_request_with_proxy(url, headers=None, max_retries=3):
    """Make a request using Oxylabs proxy with retry logic."""
    if headers is None:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
                          'AppleWebKit/537.36 (KHTML, like Gecko) '
                          'Chrome/91.0.4472.124 Safari/537.36'
        }
    
    for attempt in range(max_retries):
        try:
            if proxy_working and proxy_config:
                proxy_dict = get_proxy_dict(proxy_config)
                logger.info(f"Attempt {attempt + 1}: Using Oxylabs proxy {proxy_config['ip']}:{proxy_config['port']}")
                response = requests.get(url, headers=headers, proxies=proxy_dict, timeout=30)
            else:
                logger.warning(f"Attempt {attempt + 1}: Proxy not available, using direct connection")
                response = requests.get(url, headers=headers, timeout=30)
            
            response.raise_for_status()
            logger.info(f"Request successful on attempt {attempt + 1}")
            return response
            
        except requests.exceptions.RequestException as e:
            logger.warning(f"Request failed on attempt {attempt + 1}: {e}")
            if attempt < max_retries - 1:
                # Wait before retry with exponential backoff
                wait_time = (2 ** attempt) + random.uniform(1, 3)
                logger.info(f"Waiting {wait_time:.1f} seconds before retry...")
                time.sleep(wait_time)
            else:
                logger.error(f"All {max_retries} attempts failed for URL: {url}")
                raise
    
    return None


def cleanup_downloads_directory():
    """Delete all files in the downloads directory."""
    try:
        files_pattern = os.path.join(DOWNLOAD_DIR, "*")
        files_to_delete = glob.glob(files_pattern)
        
        if not files_to_delete:
            logger.info("Downloads directory is already empty")
            return
        
        deleted_count = 0
        total_size = 0
        
        for file_path in files_to_delete:
            try:
                if os.path.isfile(file_path):
                    file_size = os.path.getsize(file_path)
                    total_size += file_size
                    os.remove(file_path)
                    deleted_count += 1
                    logger.info(f"Deleted: {os.path.basename(file_path)} ({file_size / (1024*1024):.2f} MB)")
                    
            except Exception as e:
                logger.error(f"Error deleting {file_path}: {e}")
        
        logger.info(f"Cleanup completed: {deleted_count} files deleted, {total_size / (1024*1024):.2f} MB freed")
        
        with open(LAST_CLEANUP_FILE, 'w') as f:
            f.write(datetime.datetime.now().strftime('%Y-%m-%d'))
        
    except Exception as e:
        logger.error(f"Error during downloads cleanup: {e}")


def should_run_daily_cleanup():
    """Check if daily cleanup should be run."""
    try:
        today = datetime.datetime.now().strftime('%Y-%m-%d')
        
        if not os.path.exists(LAST_CLEANUP_FILE):
            return True
        
        with open(LAST_CLEANUP_FILE, 'r') as f:
            last_cleanup_date = f.read().strip()
        
        return last_cleanup_date != today
        
    except Exception as e:
        logger.error(f"Error checking cleanup date: {e}")
        return True


def create_nextcloud_folder(folder_path):
    """Create a folder in Nextcloud using WebDAV."""
    try:
        webdav_url = f"{NEXTCLOUD_URL}/remote.php/dav/files/{NEXTCLOUD_USERNAME}/{folder_path}"
        
        response = requests.request(
            'MKCOL',
            webdav_url,
            auth=(NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN),
            timeout=30
        )
        
        # 201 = created, 405 = already exists
        if response.status_code in [201, 405]:
            logger.info(f"Nextcloud folder ready: {folder_path}")
            return True
        else:
            logger.error(f"Failed to create Nextcloud folder: {response.status_code} - {response.text}")
            return False
            
    except Exception as e:
        logger.error(f"Error creating Nextcloud folder: {e}")
        return False


def upload_to_nextcloud(file_path, nextcloud_path):
    """Upload a file to Nextcloud using WebDAV."""
    try:
        webdav_url = f"{NEXTCLOUD_URL}/remote.php/dav/files/{NEXTCLOUD_USERNAME}/{nextcloud_path}"
        
        with open(file_path, 'rb') as f:
            response = requests.put(
                webdav_url,
                data=f,
                auth=(NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN),
                timeout=300  # 5 minutes for file upload
            )
        
        if response.status_code in [201, 204]:
            logger.info(f"Successfully uploaded to Nextcloud: {nextcloud_path}")
            return True
        else:
            logger.error(f"Failed to upload to Nextcloud: {response.status_code} - {response.text}")
            return False
            
    except Exception as e:
        logger.error(f"Error uploading to Nextcloud: {e}")
        return False


def get_nextcloud_share_link(nextcloud_path):
    """Create a public share link for a Nextcloud file with edit permissions."""
    try:
        # Create share via OCS API
        share_url = f"{NEXTCLOUD_URL}/ocs/v2.php/apps/files_sharing/api/v1/shares"
        
        data = {
            'path': f"/{nextcloud_path}",
            'shareType': 3,  # Public link
            'permissions': 15  # Full permissions (read + write + create + delete + share)
        }
        
        headers = {
            'OCS-APIRequest': 'true',
            'Content-Type': 'application/x-www-form-urlencoded'
        }
        
        response = requests.post(
            share_url,
            data=data,
            headers=headers,
            auth=(NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN),
            timeout=30
        )
        
        if response.status_code == 200:
            # Parse XML response to get share URL
            from xml.etree import ElementTree as ET
            root = ET.fromstring(response.text)
            
            url_element = root.find('.//url')
            if url_element is not None:
                share_link = url_element.text
                logger.info(f"Created Nextcloud editable share link: {share_link}")
                return share_link
            else:
                logger.error("Could not find URL in share response")
                return None
        else:
            logger.error(f"Failed to create share: {response.status_code} - {response.text}")
            return None
            
    except Exception as e:
        logger.error(f"Error creating Nextcloud share link: {e}")
        return None


def get_clean_committee_mapping():
    """Return mapping of committee names for clean folder/file naming."""
    return {
        'interim': {
            'capitol buildings planning': 'CBPC',
            'capitol security': 'Capitol Security',
            'courts corrections justice': 'CCJ',
            'courts, corrections & justice': 'CCJ',
            'economic rural development': 'ERDPC',
            'economic & rural development': 'ERDPC',
            'facilities review': 'Facilities Review',
            'federal funding stabilization': 'FFSS',
            'indian affairs': 'IAC',
            'interim legislative ethics': 'ILEC',
            'legislative ethics': 'ILEC',
            'investments pensions oversight': 'IPOC',
            'investments & pensions oversight': 'IPOC',
            'land grant': 'LGC',
            'legislative council': 'LCS',
            'legislative education study': 'LESC',
            'legislative finance': 'LFC',
            'legislative interim committee working group': 'LICW',
            'legislative health human services': 'LHHS',
	    'legislative health and human services': 'LHHS',
            'legislative health & human services': 'LHHS',
            'military veterans affairs': 'MVAC',
            'military & veterans\' affairs': 'MVAC',
            'mortgage finance authority act oversight': 'MFA',
            'mortgage finance authority oversight': 'MFA',
            'new mexico finance authority oversight': 'NMFA',
            'public school capital outlay oversight': 'PSCO',
            'public school capital outlay council': 'PSCO',
            'public school capital outlay oversight task': 'PSCO',
            'radioactive hazardous materials': 'RHMC',
            'radioactive & hazardous materials': 'RHMC',
            'revenue stabilization tax policy': 'RSTP',
            'revenue stabilization & tax policy': 'RSTP',
            'science technology telecommunications': 'STTC',
            'science, technology & telecommunications': 'STTC',
            'tobacco settlement revenue oversight': 'TSOC',
            'transportation infrastructure revenue': 'TIRS',
            'transportation infrastructure revenue subcommittee': 'TIRS',
            'water natural resources': 'WNRC',
            'water & natural resources': 'WNRC'
        },
        'house': {
            'agriculture acequias water resources': 'HAAWC',
            'agriculture, acequias and water resources': 'HAAWC',
            'appropriations finance': 'HAFC',
            'appropriations & finance': 'HAFC',
            'commerce economic development': 'HCEDC',
            'commerce & economic development committee': 'HCEDC',
            'consumer public affairs': 'HCPAC',
            'consumer & public affairs': 'HCPAC',
            'education': 'HEC',
            'energy environment natural resources': 'HENRC',
            'energy, environment & natural resources': 'HENRC',
            'government elections indian affairs': 'HGEIC',
            'government, elections & indian affairs': 'HGEIC',
            'health human services': 'HHHC',
            'health & human services': 'HHHC',
            'judiciary': 'HJC',
            'labor veterans military affairs': 'HLVMC',
            'labor, veterans\' and military affairs committee': 'HLVMC',
            'rules order business': 'HXRC',
            'rules & order of business': 'HXRC',
            'rural development land grants cultural affairs': 'HRDLC',
            'rural development, land grants and cultural affairs': 'HRDLC',
            'taxation revenue': 'HTRC',
            'taxation & revenue': 'HTRC',
            'transportation public works capital improvements': 'HTPWC',
            'transportation, public works & capital improvements': 'HTPWC'
        },
        'senate': {
            'committees committee': 'SXCC',
            'committees\' committee': 'SXCC',
            'conservation': 'SCONC',
            'education': 'SEC',
            'finance': 'SFC',
            'health public affairs': 'SHPAC',
            'health & public affairs': 'SHPAC',
            'indian rural cultural affairs': 'SIRC',
            'indian, rural & cultural affairs': 'SIRC',
            'judiciary': 'SJC',
            'rules': 'SRC',
            'tax business transportation': 'STBTC',
            'tax, business & transportation': 'STBTC'
        }
    }


def parse_meeting_title(title):
    """Parse meeting title to determine folder structure and clean title."""
    try:
        # Clean up the title first - remove extra whitespace and newlines
        clean_title = re.sub(r'\s+', ' ', title).strip()
        
        # Split on the first " - " to get prefix and committee name
        if ' - ' in clean_title:
            prefix, committee_text = clean_title.split(' - ', 1)
            prefix = prefix.strip().upper()
            committee_text = committee_text.strip()
            
            # Get committee mapping
            committee_mapping = get_clean_committee_mapping()
            
            # Clean committee text for matching (remove special chars, normalize)
            committee_for_matching = re.sub(r'[^\w\s]', ' ', committee_text.lower())
            committee_for_matching = re.sub(r'\s+', ' ', committee_for_matching).strip()
            
            # Try to find a matching committee
            clean_committee_name = None
            committee_type = None
            
            if prefix == 'IC':
                # Interim Committee
                committee_type = 'interim'
                for key, value in committee_mapping['interim'].items():
                    if key in committee_for_matching or committee_for_matching in key:
                        clean_committee_name = value
                        break
                
                if not clean_committee_name:
                    # Fallback - try partial matches
                    for key, value in committee_mapping['interim'].items():
                        key_words = set(key.split())
                        text_words = set(committee_for_matching.split())
                        if len(key_words.intersection(text_words)) >= 2:  # At least 2 words match
                            clean_committee_name = value
                            break
                
                if clean_committee_name:
                    return {
                        'type': 'interim',
                        'folder_path': f"Legislative Transcription/Interim/{clean_committee_name}",
                        'committee_name': clean_committee_name,
                        'chamber': None,
                        'clean_title': f"IC - {clean_committee_name}"
                    }
                    
            elif prefix in ['HOUSE', 'SENATE']:
                # Legislative Session Committee
                chamber = prefix.title()
                committee_type = chamber.lower()
                
                for key, value in committee_mapping[committee_type].items():
                    if key in committee_for_matching or committee_for_matching in key:
                        clean_committee_name = value
                        break
                
                if not clean_committee_name:
                    # Fallback - try partial matches
                    for key, value in committee_mapping[committee_type].items():
                        key_words = set(key.split())
                        text_words = set(committee_for_matching.split())
                        if len(key_words.intersection(text_words)) >= 2:  # At least 2 words match
                            clean_committee_name = value
                            break
                
                if clean_committee_name:
                    return {
                        'type': 'session',
                        'folder_path': f"Legislative Transcription/LegSession/{chamber}/{clean_committee_name}",
                        'committee_name': clean_committee_name,
                        'chamber': chamber,
                        'clean_title': f"{chamber} - {clean_committee_name}"
                    }
            
            # If no clean match found, log and use fallback
            logger.warning(f"Could not match committee '{committee_text}' to known committees")
            fallback_name = committee_text[:50]  # Limit length
            return {
                'type': 'unknown',
                'folder_path': f"Legislative Transcription/Other/{fallback_name}",
                'committee_name': fallback_name,
                'chamber': None,
                'clean_title': f"{prefix} - {fallback_name}"
            }
        else:
            # No " - " separator found, use default structure
            logger.warning(f"Could not parse meeting title format: {clean_title[:50]}...")
            clean_committee = re.sub(r'[^\w\s-]', '', clean_title)
            clean_committee = re.sub(r'\s+', ' ', clean_committee).strip()[:50]
            
            return {
                'type': 'unknown',
                'folder_path': f"Legislative Transcription/Other/{clean_committee or 'Unknown'}",
                'committee_name': clean_committee or 'Unknown Committee',
                'chamber': None,
                'clean_title': clean_title
            }
            
    except Exception as e:
        logger.error(f"Error parsing meeting title: {e}")
        return {
            'type': 'error',
            'folder_path': 'Legislative Transcription/Other/Parse_Error',
            'committee_name': 'Unknown Committee',
            'chamber': None,
            'clean_title': title
        }


def create_folder_hierarchy(base_folder_path, date_folder):
    """Create the full folder hierarchy including date subfolder."""
    try:
        # Create base folder structure first
        folder_parts = base_folder_path.split('/')
        current_path = ""
        
        for part in folder_parts:
            if current_path:
                current_path += f"/{part}"
            else:
                current_path = part
            
            success = create_nextcloud_folder(current_path)
            if not success:
                logger.warning(f"Could not create folder: {current_path}")
        
        # Create date subfolder
        full_path = f"{base_folder_path}/{date_folder}"
        success = create_nextcloud_folder(full_path)
        
        return success, full_path
        
    except Exception as e:
        logger.error(f"Error creating folder hierarchy: {e}")
        return False, None
    """Parse meeting title to determine folder structure and clean title."""
    try:
        # Clean up the title first - remove extra whitespace and newlines
        clean_title = re.sub(r'\s+', ' ', title).strip()
        
        # Split on the first " - " to get prefix and committee name
        if ' - ' in clean_title:
            prefix, committee_name = clean_title.split(' - ', 1)
            prefix = prefix.strip().upper()
            committee_name = committee_name.strip()
            
            # Clean committee name for folder use (remove special chars, keep spaces)
            clean_committee = re.sub(r'[^\w\s-]', '', committee_name)
            clean_committee = re.sub(r'\s+', ' ', clean_committee).strip()
            
            if prefix == 'IC':
                # Interim Committee
                folder_type = 'Interim'
                chamber = None
                return {
                    'type': 'interim',
                    'folder_path': f"Legislative Transcription/Interim/{clean_committee}",
                    'committee_name': committee_name,
                    'chamber': None,
                    'clean_title': clean_title
                }
            elif prefix in ['HOUSE', 'SENATE']:
                # Legislative Session Committee
                folder_type = 'LegSession'
                chamber = prefix.title()  # House or Senate
                return {
                    'type': 'session',
                    'folder_path': f"Legislative Transcription/LegSession/{chamber}/{clean_committee}",
                    'committee_name': committee_name,
                    'chamber': chamber,
                    'clean_title': clean_title
                }
            else:
                # Unknown prefix, fall back to original structure
                logger.warning(f"Unknown meeting prefix '{prefix}', using default folder structure")
                return {
                    'type': 'unknown',
                    'folder_path': f"Legislative Transcription/Other/{clean_committee or 'Unknown'}",
                    'committee_name': committee_name or 'Unknown Committee',
                    'chamber': None,
                    'clean_title': clean_title
                }
        else:
            # No " - " separator found, use default structure
            logger.warning(f"Could not parse meeting title format, using default folder structure: {clean_title[:50]}...")
            clean_committee = re.sub(r'[^\w\s-]', '', clean_title)
            clean_committee = re.sub(r'\s+', ' ', clean_committee).strip()[:50]  # Limit length
            
            return {
                'type': 'unknown',
                'folder_path': f"Legislative Transcription/Other/{clean_committee or 'Unknown'}",
                'committee_name': clean_committee or 'Unknown Committee',
                'chamber': None,
                'clean_title': clean_title
            }
            
    except Exception as e:
        logger.error(f"Error parsing meeting title: {e}")
        return {
            'type': 'error',
            'folder_path': 'Legislative Transcription/Other/Parse_Error',
            'committee_name': 'Unknown Committee',
            'chamber': None,
            'clean_title': title
        }


def create_folder_hierarchy(base_folder_path, date_folder):
    """Create the full folder hierarchy including date subfolder."""
    try:
        # Create base folder structure first
        folder_parts = base_folder_path.split('/')
        current_path = ""
        
        for part in folder_parts:
            if current_path:
                current_path += f"/{part}"
            else:
                current_path = part
            
            success = create_nextcloud_folder(current_path)
            if not success:
                logger.warning(f"Could not create folder: {current_path}")
        
        # Create date subfolder
        full_path = f"{base_folder_path}/{date_folder}"
        success = create_nextcloud_folder(full_path)
        
        return success, full_path
        
    except Exception as e:
        logger.error(f"Error creating folder hierarchy: {e}")
        return False, None


def check_nextcloud_file_exists(nextcloud_path):
    """Check if a file exists in Nextcloud using WebDAV."""
    try:
        webdav_url = f"{NEXTCLOUD_URL}/remote.php/dav/files/{NEXTCLOUD_USERNAME}/{nextcloud_path}"
        
        response = requests.head(
            webdav_url,
            auth=(NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN),
            timeout=30
        )
        
        # 200 = file exists, 404 = file doesn't exist
        return response.status_code == 200
        
    except Exception as e:
        logger.error(f"Error checking if Nextcloud file exists: {e}")
        return False


def generate_unique_filename(base_filename, folder_path):
    """Generate a unique filename if conflicts exist in Nextcloud."""
    try:
        # Split filename and extension
        name_part, ext = os.path.splitext(base_filename)
        counter = 1
        
        # Check if base filename exists
        test_path = f"{folder_path}/{base_filename}"
        if not check_nextcloud_file_exists(test_path):
            return base_filename
        
        # Try numbered versions until we find one that doesn't exist
        while counter <= 50:  # Reasonable limit to prevent infinite loops
            new_filename = f"{name_part}_v{counter}{ext}"
            test_path = f"{folder_path}/{new_filename}"
            
            if not check_nextcloud_file_exists(test_path):
                logger.info(f"Generated unique filename: {new_filename}")
                return new_filename
            
            counter += 1
        
        # If we still have conflicts after 50 tries, add timestamp
        timestamp = datetime.datetime.now().strftime('%H%M%S')
        unique_filename = f"{name_part}_{timestamp}{ext}"
        logger.warning(f"Used timestamp for unique filename: {unique_filename}")
        return unique_filename
        
    except Exception as e:
        logger.error(f"Error generating unique filename: {e}")
        # Fallback to timestamp-based name
        timestamp = datetime.datetime.now().strftime('%H%M%S')
        name_part, ext = os.path.splitext(base_filename)
        return f"{name_part}_{timestamp}{ext}"


def save_transcript_to_nextcloud(transcript_data, title, entry_url):
    """Save transcript as DOCX to Nextcloud with structured folder organization."""
    try:
        if not all([NEXTCLOUD_URL, NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN]):
            logger.error("Nextcloud configuration missing")
            return None
        
        # Parse the meeting title to determine folder structure
        meeting_info = parse_meeting_title(title)
        
        # Create date folder (YYYY-MM-DD format)
        date_folder = datetime.datetime.now().strftime('%Y-%m-%d')
        
        # Create the full folder hierarchy
        success, full_folder_path = create_folder_hierarchy(meeting_info['folder_path'], date_folder)
        
        if not success:
            logger.error("Failed to create folder hierarchy, falling back to simple structure")
            # Fallback to simple structure
            simple_folder = f"Legislative Transcription/{date_folder}"
            create_nextcloud_folder("Legislative Transcription")
            create_nextcloud_folder(simple_folder)
            full_folder_path = simple_folder
        
        # Create base filename (YYYYMMDD-[Meeting Title].docx)
        date_prefix = datetime.datetime.now().strftime('%Y%m%d')
        
        # Use clean title from parsing
        clean_title = meeting_info['clean_title']
        # Further clean for filename (remove special chars, limit length)
        filename_title = re.sub(r'[^\w\s-]', '', clean_title)[:80]  # Limit length for filename
        filename_title = re.sub(r'\s+', ' ', filename_title).strip()
        
        base_filename = f"{date_prefix}-{filename_title}.docx"
        
        # Generate unique filename if conflicts exist
        unique_filename = generate_unique_filename(base_filename, full_folder_path)
        
        # Create the DOCX document with enhanced metadata
        doc = create_docx_document_with_metadata(transcript_data, title, entry_url, meeting_info)
        if not doc:
            logger.error("Failed to create DOCX document")
            return None
        
        # Save to temporary file
        temp_path = os.path.join(DOWNLOAD_DIR, unique_filename)
        doc.save(temp_path)
        
        # Upload to Nextcloud
        nextcloud_file_path = f"{full_folder_path}/{unique_filename}"
        
        if upload_to_nextcloud(temp_path, nextcloud_file_path):
            logger.info(f"Successfully saved transcript to: {nextcloud_file_path}")
            
            # Create share link
            share_link = get_nextcloud_share_link(nextcloud_file_path)
            
            # Clean up temporary file
            try:
                os.remove(temp_path)
            except Exception as e:
                logger.warning(f"Could not remove temporary file: {e}")
            
            return {
                'nextcloud_path': nextcloud_file_path,
                'share_link': share_link,
                'filename': unique_filename,
                'meeting_info': meeting_info,
                'folder_path': full_folder_path
            }
        else:
            logger.error("Failed to upload to Nextcloud")
            return None
            
    except Exception as e:
        logger.error(f"Error saving transcript to Nextcloud: {e}")
        return None


def create_docx_document_with_metadata(transcript_data, title, entry_url, meeting_info):
    """Create a DOCX document with enhanced meeting metadata."""
    try:
        # Create a new document
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add meeting information section
        doc.add_heading('Meeting Information', level=1)
        
        # Committee information
        if meeting_info['type'] == 'interim':
            info_paragraph = doc.add_paragraph()
            info_paragraph.add_run('Meeting Type: ').bold = True
            info_paragraph.add_run('Interim Committee')
            
        elif meeting_info['type'] == 'session':
            info_paragraph = doc.add_paragraph()
            info_paragraph.add_run('Meeting Type: ').bold = True
            info_paragraph.add_run('Legislative Session Committee')
            
            info_paragraph = doc.add_paragraph()
            info_paragraph.add_run('Chamber: ').bold = True
            info_paragraph.add_run(meeting_info['chamber'])
        
        info_paragraph = doc.add_paragraph()
        info_paragraph.add_run('Committee: ').bold = True
        info_paragraph.add_run(meeting_info['committee_name'])
        
        # Standard metadata
        info_paragraph = doc.add_paragraph()
        info_paragraph.add_run('Source URL: ').bold = True
        info_paragraph.add_run(entry_url)
        
        info_paragraph = doc.add_paragraph()
        info_paragraph.add_run('Transcribed by: ').bold = True
        info_paragraph.add_run('LWE.Vote')
        
        info_paragraph = doc.add_paragraph()
        info_paragraph.add_run('Transcribed on: ').bold = True
        info_paragraph.add_run(datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        
        if transcript_data:
            info_paragraph = doc.add_paragraph()
            info_paragraph.add_run('Audio duration: ').bold = True
            info_paragraph.add_run(f"{transcript_data.get('audio_duration', 'Unknown')} seconds")
            
            info_paragraph = doc.add_paragraph()
            info_paragraph.add_run('Confidence score: ').bold = True
            info_paragraph.add_run(f"{transcript_data.get('confidence', 'Unknown')}")
        
        # Add separator
        doc.add_paragraph('_' * 80)
        
        # Add transcript section
        doc.add_heading('Transcript', level=1)
        
        # Format and add transcript content
        if transcript_data and "words" in transcript_data:
            current_speaker = None
            current_paragraph = None
            
            for word in transcript_data["words"]:
                speaker = word.get("speaker")
                
                if speaker != current_speaker:
                    # New speaker, create new paragraph
                    current_paragraph = doc.add_paragraph()
                    current_speaker = speaker
                    
                    if speaker is not None:
                        speaker_run = current_paragraph.add_run(f"Speaker {speaker}: ")
                        speaker_run.bold = True
                        current_paragraph.add_run(f"{word['text']} ")
                    else:
                        current_paragraph.add_run(f"{word['text']} ")
                else:
                    # Same speaker, continue paragraph
                    if current_paragraph:
                        current_paragraph.add_run(f"{word['text']} ")
        
        elif transcript_data and transcript_data.get("text"):
            # Fallback to plain text if words are not available
            doc.add_paragraph(transcript_data["text"])
        else:
            doc.add_paragraph("No transcription text available.")
        
        return doc
        
    except Exception as e:
        logger.error(f"Error creating DOCX document with metadata: {e}")
        return None


def get_current_entries():
    """Fetch and parse the current entries from the website using proxy."""
    try:
        # Use proxy-enabled request
        response = make_request_with_proxy(URL)
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        entries = []
        
        for entry_element in soup.select('table'):
            entry_text = entry_element.get_text().strip()
            
            parent_a = entry_element.find_parent('a')
            entry_link = None
            if parent_a and 'href' in parent_a.attrs:
                entry_link = parent_a['href']
                if entry_link and not entry_link.startswith('http'):
                    base_url = '/'.join(URL.split('/')[:3])
                    entry_link = f"{base_url}{entry_link if entry_link.startswith('/') else '/' + entry_link}"
            
            if entry_text:
                entries.append({
                    'text': entry_text,
                    'link': entry_link,
                    'timestamp': datetime.datetime.now().isoformat()
                })
        
        return entries
    
    except Exception as e:
        logger.error(f"Error fetching entries: {e}")
        return []


def read_stored_entries():
    """Read the previously stored entries."""
    if not os.path.exists(ENTRIES_FILE):
        return []
    
    try:
        with open(ENTRIES_FILE, 'r') as f:
            return json.load(f)
    except json.JSONDecodeError:
        logger.error(f"Error parsing {ENTRIES_FILE}, starting with empty entries")
        return []


def write_entries(entries):
    """Write entries to the storage file."""
    with open(ENTRIES_FILE, 'w') as f:
        json.dump(entries, f, indent=2)


def read_processed_entries():
    """Read the previously processed entries with error handling for format changes."""
    if not os.path.exists(PROCESSED_ENTRIES_FILE):
        return []
    
    try:
        with open(PROCESSED_ENTRIES_FILE, 'r') as f:
            content = f.read().strip()
            if not content:
                return []
            
            entries = json.loads(content)
            
            # Validate that entries is a list
            if not isinstance(entries, list):
                logger.error(f"Processed entries file contains invalid format (not a list), resetting...")
                return []
            
            # Check if entries have expected structure and clean up if needed
            valid_entries = []
            for entry in entries:
                if isinstance(entry, dict) and 'text' in entry:
                    # Convert old Google Docs format to new Nextcloud format if needed
                    if 'transcription' in entry:
                        transcription = entry['transcription']
                        if isinstance(transcription, dict):
                            # Update old google_doc_url to nextcloud_result format
                            if 'google_doc_url' in transcription and 'nextcloud_result' not in transcription:
                                transcription['nextcloud_result'] = {
                                    'legacy_google_doc': transcription.pop('google_doc_url'),
                                    'migrated_to_nextcloud': True
                                }
                    
                    valid_entries.append(entry)
                else:
                    logger.warning(f"Skipping invalid entry format: {entry}")
            
            return valid_entries
            
    except json.JSONDecodeError as e:
        logger.error(f"Error parsing {PROCESSED_ENTRIES_FILE}: {e}")
        logger.info(f"Backing up corrupted file and starting fresh...")
        
        # Backup the corrupted file
        backup_file = f"{PROCESSED_ENTRIES_FILE}.backup.{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
        try:
            os.rename(PROCESSED_ENTRIES_FILE, backup_file)
            logger.info(f"Backed up corrupted file to: {backup_file}")
        except Exception as backup_error:
            logger.error(f"Could not backup corrupted file: {backup_error}")
        
        return []
    except Exception as e:
        logger.error(f"Unexpected error reading processed entries: {e}")
        return []


def write_processed_entry(entry):
    """Add an entry to the processed entries file."""
    processed = read_processed_entries()
    processed.append(entry)
    
    with open(PROCESSED_ENTRIES_FILE, 'w') as f:
        json.dump(processed, f, indent=2)


def send_notification(new_entries, transcript_results=None):
    """Send email notification about new entries with Nextcloud links."""
    if not all([EMAIL_USER, EMAIL_PASSWORD]):
        logger.warning("Email configuration missing. Skipping notification.")
        return
    
    try:
        for recipient in EMAIL_RECIPIENTS:
            msg = MIMEMultipart()
            msg['From'] = EMAIL_USER
            msg['To'] = recipient
            msg['Subject'] = f"New Legislature Entries - {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"
            
            body = "New legislature meeting entries have been detected and transcribed:\n\n"
            
            for i, entry in enumerate(new_entries):
                body += f"Meeting: {entry['text']}\n"
                if entry['link']:
                    body += f"Source: {entry['link']}\n"
                
                if transcript_results and i < len(transcript_results) and transcript_results[i]:
                    result = transcript_results[i]
                    if result.get('share_link'):
                        body += f"Transcript: {result['share_link']}\n"
                    elif result.get('nextcloud_path'):
                        body += f"Transcript saved to: {result['nextcloud_path']}\n"
                    
                    if result.get('filename'):
                        body += f"Filename: {result['filename']}\n"
                
                body += "\n" + "-"*50 + "\n\n"
            
            body += f"View all meetings at: {URL}\n\n"
            body += "This notification was sent by the Legislature Monitoring System."
            
            msg.attach(MIMEText(body, 'plain'))
            
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login(EMAIL_USER, EMAIL_PASSWORD)
            server.send_message(msg)
            server.quit()
            
            logger.info(f"Notification email sent successfully to {recipient}")
            
    except Exception as e:
        logger.error(f"Error sending notification: {e}")


def extract_hls_stream_url(entry_url):
    """Extract the HLS stream URL from the entry page using proxy."""
    try:
        logger.info(f"Accessing entry page to extract HLS stream: {entry_url}")
        
        # Use proxy-enabled request
        response = make_request_with_proxy(entry_url)

        # First try to find the availableStreams JSON data
        available_streams_pattern = r'var\s+availableStreams\s*=\s*(\[.*?\]);'
        available_streams_match = re.search(available_streams_pattern,
                                            response.text,
                                            re.DOTALL)
        if available_streams_match:
            streams_json_str = available_streams_match.group(1)
            try:
                streams_data = json.loads(streams_json_str)
                logger.info(f"Found {len(streams_data)} streams in availableStreams JSON")
                
                # Look for the best stream (non-live, enabled, with URL ending in .m3u8)
                best_stream = None
                for stream in streams_data:
                    url = stream.get('Url', '')
                    is_live = stream.get('IsLive', True)
                    enabled = stream.get('Enabled', False)
                    
                    logger.info(f"Stream: URL={url[:80]}{'...' if len(url) > 80 else ''}, IsLive={is_live}, Enabled={enabled}")
                    
                    if url.endswith('.m3u8') and enabled and not is_live:
                        # This is a good recorded stream
                        best_stream = stream
                        break
                    elif url.endswith('.m3u8') and enabled:
                        # Fallback to any enabled stream
                        if best_stream is None:
                            best_stream = stream
                
                if best_stream:
                    # Clean up the URL - replace escaped forward slashes but preserve URL encoding
                    hls_url = best_stream['Url'].replace('\\/', '/')
                    logger.info(f"Selected stream: IsLive={best_stream.get('IsLive')}, Duration={best_stream.get('Duration')}s")
                    logger.info(f"Extracted HLS URL from JSON: {hls_url}")
                    return hls_url
                else:
                    logger.warning("No suitable streams found in availableStreams JSON")
                    
            except json.JSONDecodeError as e:
                logger.error(f"Could not parse availableStreams JSON: {e}")
                logger.debug(f"JSON string was: {streams_json_str[:200]}...")

        # Enhanced fallback patterns - look for VOD URLs first
        vod_patterns = [
            # Pattern for VOD URLs with _definst_ and encoded title (most specific)
            r'(https?://[^"\'\s]+vod-2/_definst_/[^"\'\s]+\.mp4/playlist\.m3u8)',
            # Pattern for any VOD URL
            r'(https?://[^"\'\s]+vod[^"\'\s]*\.m3u8)',
            # General m3u8 pattern but prioritize those with dates/times
            r'(https?://[^"\'\s]+/\d{4}/\d{2}/\d{2}/[^"\'\s]+\.m3u8)',
        ]
        
        for i, pattern in enumerate(vod_patterns):
            matches = re.findall(pattern, response.text, re.IGNORECASE)
            if matches:
                logger.info(f"Found {len(matches)} matches with pattern {i+1}")
                # Filter out live streams if we have multiple matches
                for match in matches:
                    if 'live' not in match.lower():
                        logger.info(f"Found VOD HLS URL via regex pattern {i+1}: {match}")
                        return match
                # If all have 'live', take the first one
                logger.info(f"Found HLS URL via regex pattern {i+1} (might be live): {matches[0]}")
                return matches[0]

        # Last resort - any m3u8 URL
        general_m3u8_pattern = r'(https?://[^"\'\s]+\.m3u8)'
        matches = re.findall(general_m3u8_pattern, response.text)
        if matches:
            logger.info(f"Found {len(matches)} m3u8 URLs as last resort")
            # Filter out known live stream patterns
            for match in matches:
                if '/live/' not in match and 'playlist.m3u8' in match:
                    logger.info(f"Found potential HLS URL via fallback: {match}")
                    return match
            
            # If no good matches, return first one
            logger.warning(f"Using first available HLS URL (may be live stream): {matches[0]}")
            return matches[0]

        logger.error("No HLS stream URLs found on the page")
        return None

    except Exception as e:
        logger.error(f"Error extracting HLS stream URL: {e}")
        return None


def download_hls_with_ytdlp(hls_url, output_file, max_retries=3):
    """Download HLS stream using yt-dlp with Oxylabs proxy."""
    try:
        output_dir = os.path.dirname(output_file)
        if output_dir:
            Path(output_dir).mkdir(exist_ok=True)
        
        logger.info(f"Downloading HLS stream with yt-dlp from {hls_url} to {output_file}")
        
        # Try multiple attempts
        for attempt in range(max_retries):
            try:
                # Configure yt-dlp options
                ydl_opts = {
                    'outtmpl': output_file.replace('.mp4', '.%(ext)s'),
                    'format': 'best[ext=mp4]/best',  # Prefer mp4, fallback to best available
                    'writesubtitles': False,
                    'writeautomaticsub': False,
                    'ignoreerrors': False,
                    'no_warnings': False,
                    'extractaudio': False,
                    'socket_timeout': 60,
                    'retries': 3,
                    'fragment_retries': 5,
                    'http_chunk_size': 10485760,  # 10MB chunks
                    'hls_prefer_native': True,    # Use native HLS extractor
                }
                
                # Add proxy configuration if available
                if proxy_working and proxy_config:
                    proxy_url = get_proxy_url(proxy_config)
                    ydl_opts['proxy'] = proxy_url
                    logger.info(f"Attempt {attempt + 1}: Using Oxylabs proxy {proxy_config['ip']}:{proxy_config['port']} with yt-dlp")
                else:
                    logger.warning(f"Attempt {attempt + 1}: Proxy not available, using direct connection")
                
                # Add user agent
                ydl_opts['http_headers'] = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                
                # Download with yt-dlp
                with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                    ydl.download([hls_url])
                
                # Check if file was created (yt-dlp might change extension)
                possible_files = [
                    output_file,
                    output_file.replace('.mp4', '.mkv'),
                    output_file.replace('.mp4', '.webm'),
                    output_file.replace('.mp4', '.m4v')
                ]
                
                downloaded_file = None
                for possible_file in possible_files:
                    if os.path.exists(possible_file) and os.path.getsize(possible_file) > 1024:
                        downloaded_file = possible_file
                        break
                
                if downloaded_file:
                    # Rename to desired output file if necessary
                    if downloaded_file != output_file:
                        os.rename(downloaded_file, output_file)
                        downloaded_file = output_file
                    
                    file_size_mb = os.path.getsize(downloaded_file) / (1024 * 1024)
                    logger.info(f"Successfully downloaded video to {downloaded_file} (size: {file_size_mb:.2f} MB)")
                    return downloaded_file
                else:
                    logger.error("Download failed - no output file found")
                    
            except Exception as e:
                logger.warning(f"yt-dlp attempt {attempt + 1} failed: {e}")
                if attempt < max_retries - 1:
                    wait_time = (2 ** attempt) + random.uniform(1, 3)
                    logger.info(f"Waiting {wait_time:.1f} seconds before retry...")
                    time.sleep(wait_time)
                continue
        
        logger.error(f"All {max_retries} yt-dlp attempts failed")
        return None
        
    except Exception as e:
        logger.error(f"Error in yt-dlp download process: {e}")
        return None


def download_with_ytdlp_fallback(hls_url, output_file):
    """Download with yt-dlp, fallback to direct if proxy fails."""
    try:
        # First try with proxy
        result = download_hls_with_ytdlp(hls_url, output_file)
        if result:
            return result
        
        # Fallback to direct connection
        logger.info("Trying yt-dlp with direct connection...")
        
        ydl_opts = {
            'outtmpl': output_file.replace('.mp4', '.%(ext)s'),
            'format': 'best[ext=mp4]/best',
            'writesubtitles': False,
            'writeautomaticsub': False,
            'ignoreerrors': False,
            'socket_timeout': 60,
            'retries': 5,
            'fragment_retries': 10,
            'http_chunk_size': 10485760,
            'hls_prefer_native': True,
            'http_headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([hls_url])
        
        # Check for downloaded file
        possible_files = [
            output_file,
            output_file.replace('.mp4', '.mkv'),
            output_file.replace('.mp4', '.webm'),
            output_file.replace('.mp4', '.m4v')
        ]
        
        for possible_file in possible_files:
            if os.path.exists(possible_file) and os.path.getsize(possible_file) > 1024:
                if possible_file != output_file:
                    os.rename(possible_file, output_file)
                
                file_size_mb = os.path.getsize(output_file) / (1024 * 1024)
                logger.info(f"Direct connection successful (size: {file_size_mb:.2f} MB)")
                return output_file
        
        logger.error("Direct connection also failed")
        return None
        
    except Exception as e:
        logger.error(f"yt-dlp fallback failed: {e}")
        return None


def download_video(entry_url):
    """Download video from the entry page using HLS stream extraction with yt-dlp."""
    try:
        hls_url = extract_hls_stream_url(entry_url)
        if not hls_url:
            logger.error("Could not find HLS stream URL on the page")
            return None
        
        video_filename = os.path.join(
            DOWNLOAD_DIR, 
            f"video_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.mp4"
        )
        
        # Use yt-dlp with fallback strategies
        return download_with_ytdlp_fallback(hls_url, video_filename)
    
    except Exception as e:
        logger.error(f"Error in video download process: {e}")
        return None


def extract_audio_from_video(video_path):
    """Extract audio from video file using ffmpeg directly (more reliable for local files)."""
    try:
        if not os.path.exists(video_path):
            logger.error(f"Video file not found: {video_path}")
            return None
        
        audio_path = video_path.replace('.mp4', '.mp3')
        
        logger.info(f"Extracting audio from {video_path} to {audio_path}")
        
        # Use ffmpeg directly for local file conversion (more reliable than yt-dlp for this)
        try:
            result = subprocess.run([
                'ffmpeg', '-i', video_path, 
                '-vn',  # No video
                '-acodec', 'libmp3lame',  # MP3 codec
                '-ab', '192k',  # 192kbps bitrate
                '-ar', '44100',  # Sample rate
                '-y',  # Overwrite output file
                audio_path
            ], check=True, capture_output=True, text=True)
            
            if os.path.exists(audio_path) and os.path.getsize(audio_path) > 1024:
                file_size_mb = os.path.getsize(audio_path) / (1024 * 1024)
                logger.info(f"Audio extracted successfully to {audio_path} (size: {file_size_mb:.2f} MB)")
                return audio_path
            else:
                logger.error("Audio extraction failed - output file not created or too small")
                return None
                
        except subprocess.CalledProcessError as e:
            logger.error(f"ffmpeg failed: {e.stderr}")
            # Try alternative ffmpeg method
            return extract_audio_with_alternative_ffmpeg(video_path)
            
        except FileNotFoundError:
            logger.error("ffmpeg not found. Cannot extract audio.")
            return None
    
    except Exception as e:
        logger.error(f"Error extracting audio: {e}")
        return None


def extract_audio_with_alternative_ffmpeg(video_path):
    """Alternative ffmpeg method with different parameters."""
    try:
        audio_path = video_path.replace('.mp4', '.mp3')
        
        logger.info(f"Trying alternative ffmpeg method...")
        
        # Alternative method with simpler parameters
        result = subprocess.run([
            'ffmpeg', '-i', video_path, 
            '-q:a', '2',  # High quality audio
            '-map', 'a',  # Map audio stream
            '-y',  # Overwrite
            audio_path
        ], check=True, capture_output=True, text=True)
        
        if os.path.exists(audio_path) and os.path.getsize(audio_path) > 1024:
            file_size_mb = os.path.getsize(audio_path) / (1024 * 1024)
            logger.info(f"Audio extracted with alternative method to {audio_path} (size: {file_size_mb:.2f} MB)")
            return audio_path
        else:
            logger.error("Alternative audio extraction failed")
            return None
            
    except Exception as e:
        logger.error(f"Alternative ffmpeg method failed: {e}")
        return None


def transcribe_with_assemblyai(audio_path):
    """Transcribe audio using AssemblyAI API."""
    if not ASSEMBLYAI_API_KEY:
        logger.error("AssemblyAI API key not found in environment variables")
        return None
    
    try:
        headers = {
            "authorization": ASSEMBLYAI_API_KEY,
            "content-type": "application/json"
        }
        
        logger.info(f"Uploading audio file to AssemblyAI: {audio_path}")
        
        with open(audio_path, "rb") as f:
            response = requests.post(
                f"{ASSEMBLYAI_BASE_URL}/upload",
                headers=headers,
                data=f
            )
            response.raise_for_status()
        
        upload_url = response.json()["upload_url"]
        logger.info(f"Audio uploaded successfully. URL: {upload_url}")
        
        logger.info("Submitting transcription request")
        response = requests.post(
            f"{ASSEMBLYAI_BASE_URL}/transcript",
            headers=headers,
            json={
                "audio_url": upload_url,
                "speaker_labels": True,
                "punctuate": True,
                "format_text": True
            }
        )
        response.raise_for_status()
        
        transcript_id = response.json()["id"]
        logger.info(f"Transcription request submitted. ID: {transcript_id}")
        
        logger.info("Waiting for transcription to complete...")
        
        while True:
            response = requests.get(
                f"{ASSEMBLYAI_BASE_URL}/transcript/{transcript_id}",
                headers=headers
            )
            response.raise_for_status()
            
            status = response.json()["status"]
            
            if status == "completed":
                logger.info("Transcription completed successfully")
                return response.json()
            elif status == "error":
                logger.error(f"Transcription failed: {response.json().get('error')}")
                return None
            
            logger.info(f"Transcription status: {status}. Waiting...")
            time.sleep(30)  # Poll every 30 seconds
    
    except Exception as e:
        logger.error(f"Error in transcription process: {e}")
        return None


def process_new_entry(entry):
    """Process a new entry: download video, transcribe, and save to Nextcloud."""
    try:
        if not entry.get('link'):
            logger.error("Entry has no link, cannot process")
            return None
        
        video_path = download_video(entry['link'])
        if not video_path:
            logger.error("Failed to download video")
            return None
        
        audio_path = extract_audio_from_video(video_path)
        if not audio_path:
            logger.error("Failed to extract audio")
            return None
        
        transcript_data = transcribe_with_assemblyai(audio_path)
        if not transcript_data:
            logger.error("Failed to transcribe audio")
            return None
        
        # Clean title for use in filename and document
        title_text = entry['text']
        clean_title = re.sub(r'\s+', ' ', title_text).strip()
        
        nextcloud_result = save_transcript_to_nextcloud(transcript_data, clean_title, entry['link'])
        
        entry['transcription'] = {
            'video_path': video_path,
            'audio_path': audio_path,
            'nextcloud_result': nextcloud_result,
            'processed_at': datetime.datetime.now().isoformat()
        }
        
        write_processed_entry(entry)
        
        return nextcloud_result
    
    except Exception as e:
        logger.error(f"Error processing entry: {e}")
        return None


def test_proxy_connectivity():
    """Test if Oxylabs proxy is working correctly."""
    logger.info("Testing Oxylabs proxy connectivity...")
    
    test_url = "http://httpbin.org/ip"  # Simple IP echo service
    
    if proxy_working and proxy_config:
        proxy_dict = get_proxy_dict(proxy_config)
        
        try:
            response = requests.get(test_url, proxies=proxy_dict, timeout=10)
            response.raise_for_status()
            ip_data = response.json()
            logger.info(f"Oxylabs proxy working - IP: {ip_data.get('origin', 'unknown')}")
        except Exception as e:
            logger.warning(f"Oxylabs proxy test failed: {e}")
    
    # Test without proxy
    try:
        response = requests.get(test_url, timeout=10)
        response.raise_for_status()
        ip_data = response.json()
        logger.info(f"Direct connection - Working (IP: {ip_data.get('origin', 'unknown')})")
    except Exception as e:
        logger.warning(f"Direct connection - Failed: {e}")


def is_test_meeting(entry_text):
    """Check if an entry is a test meeting that should be skipped."""
    # Convert to lowercase for case-insensitive matching
    text_lower = entry_text.lower()
    
    # Check for "test meeting" anywhere in the text
    if "test meeting" in text_lower:
        return True
    
    return False


def check_for_new_entries():
    """Check for new entries and process them if found."""
    # Check if we need to test proxy
    if should_update_proxy_list():
        logger.info("Testing Oxylabs proxy...")
        test_oxylabs_proxy()
    
    # Check if we need to run daily cleanup
    if should_run_daily_cleanup():
        logger.info("Running daily cleanup of downloads directory...")
        cleanup_downloads_directory()
    
    logger.info(f"Checking for new entries at {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    current_entries = get_current_entries()
    stored_entries = read_stored_entries()
    
    stored_texts = [entry['text'] for entry in stored_entries]
    new_entries = [entry for entry in current_entries if entry['text'] not in stored_texts]
    
    # Filter out test meetings
    filtered_new_entries = []
    skipped_test_meetings = 0
    
    for entry in new_entries:
        if is_test_meeting(entry['text']):
            logger.info(f"Skipping test meeting: {entry['text'][:50]}...")
            skipped_test_meetings += 1
        else:
            filtered_new_entries.append(entry)
    
    if skipped_test_meetings > 0:
        logger.info(f"Skipped {skipped_test_meetings} test meeting(s)")
    
    if filtered_new_entries:
        logger.info(f"Found {len(filtered_new_entries)} new entries to process!")
        
        transcript_results = []
        for entry in filtered_new_entries:
            logger.info(f"Processing entry: {entry['text'][:50]}...")
            result = process_new_entry(entry)
            transcript_results.append(result)
        
        send_notification(filtered_new_entries, transcript_results)
        write_entries(current_entries)  # Still save all entries (including test meetings) to avoid reprocessing
    else:
        if new_entries:
            logger.info("All new entries were test meetings - no processing needed")
        else:
            logger.info("No new entries found")


def main():
    """Main function to schedule and run the monitoring."""
    logger.info("Starting Legislature page monitoring with yt-dlp, AssemblyAI, Nextcloud integration, and Oxylabs Proxy...")
    
    # Check Nextcloud configuration
    if not all([NEXTCLOUD_URL, NEXTCLOUD_USERNAME, NEXTCLOUD_TOKEN]):
        logger.error("Nextcloud configuration missing. Please check NEXTCLOUD_URL, NEXTCLOUD_USERNAME, and NEXTCLOUD_TOKEN in .env file")
        return
    else:
        logger.info(f"Nextcloud configured: {NEXTCLOUD_URL}")
    
    # Initialize Oxylabs proxy system
    logger.info("Initializing Oxylabs proxy system...")
    if not test_oxylabs_proxy():
        logger.warning("Failed to connect to Oxylabs proxy. Script will attempt to run without proxy.")
    else:
        logger.info("Oxylabs proxy system initialized successfully")
        # Test proxy connectivity
        test_proxy_connectivity()
    
    # Check for required dependencies
    try:
        import yt_dlp
        logger.info("yt-dlp is available")
    except ImportError:
        logger.error("yt-dlp not found. Please install with: pip install yt-dlp")
        return
    
    try:
        from docx import Document
        logger.info("python-docx is available")
    except ImportError:
        logger.error("python-docx not found. Please install with: pip install python-docx")
        return
    
    # Check for ffmpeg (optional, for fallback)
    try:
        subprocess.run(['ffmpeg', '-version'], check=True, capture_output=True)
        logger.info("ffmpeg is available (for fallback audio extraction)")
    except (subprocess.SubprocessError, FileNotFoundError):
        logger.warning("ffmpeg not found. Audio extraction will rely on yt-dlp only.")
    
    if not ASSEMBLYAI_API_KEY:
        logger.warning("AssemblyAI API key not found. Transcription will not work.")
    
    logger.info("All dependencies checked. Starting monitoring system...")
    
    # Run immediately on startup
    check_for_new_entries()
    
    # Schedule checks every few hours, Monday through Friday only (GMT times for MST = GMT-7)
    gmt_slots = ["16:00", "18:00", "20:00", "22:00", "00:00", "02:00"]

    for t in gmt_slots:
        schedule.every().monday.at(t).do(check_for_new_entries)
        schedule.every().tuesday.at(t).do(check_for_new_entries)
        schedule.every().wednesday.at(t).do(check_for_new_entries)
        schedule.every().thursday.at(t).do(check_for_new_entries)
        schedule.every().friday.at(t).do(check_for_new_entries)
    
    logger.info(f"Scheduled monitoring Monday-Friday at GMT times: {', '.join(gmt_slots)}")
    logger.info("Legislature monitoring system is now running (weekdays only)...")
    
    # Keep the script running
    while True:
        schedule.run_pending()
        time.sleep(60)


if __name__ == "__main__":
    main()
